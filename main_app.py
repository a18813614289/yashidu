import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import os
import sys
import traceback
import threading
import time
import re

# 确保中文正常显示
import locale
try:
    locale.setlocale(locale.LC_ALL, 'zh_CN.UTF-8')
except:
    try:
        locale.setlocale(locale.LC_ALL, '')
    except:
        pass

# 导入必要的库
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import lxml.etree
import win32com.client

class ExcelToWordApp:
    def __init__(self, root):
        self.root = root
        self.root.title("压实度文档处理工具")
        self.root.geometry("800x600")
        self.root.resizable(True, True)
        
        # 设置中文字体
        self.style = ttk.Style()
        if sys.platform == 'win32':
            self.style.configure(
                'TLabel',
                font=('Microsoft YaHei UI', 10)
            )
            self.style.configure(
                'TButton',
                font=('Microsoft YaHei UI', 10)
            )
        
        # 变量定义
        self.excel_path = tk.StringVar()
        self.word_path = tk.StringVar()
        self.output_path = tk.StringVar()
        self.copy_count = tk.StringVar(value="1")
        self.status_var = tk.StringVar(value="就绪")
        self.progress_var = tk.DoubleVar()
        self.is_processing = False
        self.process_thread = None
        
        # 处理选项
        self.option_process_gcm3 = tk.BooleanVar(value=True)
        self.option_unify_fonts = tk.BooleanVar(value=True)
        self.option_delete_empty_rows = tk.BooleanVar(value=True)
        
        # 创建界面
        self.create_widgets()
        
        # 设置工作目录
        self.set_working_directory()
    
    def create_widgets(self):
        # 创建主框架
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 文件选择区域
        file_frame = ttk.LabelFrame(main_frame, text="文件选择", padding="10")
        file_frame.pack(fill=tk.X, pady=(0, 10))
        
        # Excel文件选择
        ttk.Label(file_frame, text="Excel文件:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        ttk.Entry(file_frame, textvariable=self.excel_path, width=50).grid(row=0, column=1, padx=5, pady=5)
        ttk.Button(file_frame, text="浏览...", command=self._browse_excel).grid(row=0, column=2, padx=5, pady=5)
        
        # Word模板选择
        ttk.Label(file_frame, text="Word模板:").grid(row=1, column=0, padx=5, pady=5, sticky="w")
        ttk.Entry(file_frame, textvariable=self.word_path, width=50).grid(row=1, column=1, padx=5, pady=5)
        ttk.Button(file_frame, text="浏览...", command=self._browse_word).grid(row=1, column=2, padx=5, pady=5)
        
        # 输出文件路径
        ttk.Label(file_frame, text="输出路径:").grid(row=2, column=0, padx=5, pady=5, sticky="w")
        ttk.Entry(file_frame, textvariable=self.output_path, width=50).grid(row=2, column=1, padx=5, pady=5)
        ttk.Button(file_frame, text="浏览...", command=self._browse_output).grid(row=2, column=2, padx=5, pady=5)
        
        # 复制次数
        ttk.Label(file_frame, text="复制次数:").grid(row=3, column=0, padx=5, pady=5, sticky="w")
        ttk.Entry(file_frame, textvariable=self.copy_count, width=10).grid(row=3, column=1, padx=5, pady=5, sticky="w")
        
        # 处理选项
        options_frame = ttk.LabelFrame(main_frame, text="处理选项", padding="10")
        options_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Checkbutton(options_frame, text="处理g/cm³单位", variable=self.option_process_gcm3).grid(row=0, column=0, padx=10, pady=5, sticky="w")
        ttk.Checkbutton(options_frame, text="统一附表标题字体", variable=self.option_unify_fonts).grid(row=0, column=1, padx=10, pady=5, sticky="w")
        ttk.Checkbutton(options_frame, text="删除空数据行", variable=self.option_delete_empty_rows).grid(row=0, column=2, padx=10, pady=5, sticky="w")
        
        # 进度条
        progress_frame = ttk.Frame(main_frame)
        progress_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.progress_bar = ttk.Progressbar(progress_frame, variable=self.progress_var, length=100)
        self.progress_bar.pack(fill=tk.X, padx=5)
        
        self.progress_label = ttk.Label(progress_frame, text="准备就绪")
        self.progress_label.pack(padx=5, pady=5)
        
        # 按钮区域
        buttons_frame = ttk.Frame(main_frame)
        buttons_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.process_button = ttk.Button(buttons_frame, text="开始处理", command=self._start_processing)
        self.process_button.pack(side=tk.LEFT, padx=5)
        
        self.cancel_button = ttk.Button(buttons_frame, text="取消", command=self._cancel_processing, state=tk.DISABLED)
        self.cancel_button.pack(side=tk.LEFT, padx=5)
        
        # 日志区域
        log_frame = ttk.LabelFrame(main_frame, text="运行日志", padding="10")
        log_frame.pack(fill=tk.BOTH, expand=True)
        
        self.log_text = scrolledtext.ScrolledText(log_frame, wrap=tk.WORD, state=tk.DISABLED, font=("SimHei", 9))
        self.log_text.pack(fill=tk.BOTH, expand=True)
        
        # 状态栏
        self.status_bar = ttk.Label(
            self.root, 
            textvariable=self.status_var, 
            relief=tk.SUNKEN, 
            anchor=tk.W, 
            font=("Microsoft YaHei UI", 9)
        )
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)
    
    def _browse_excel(self):
        file_path = filedialog.askopenfilename(
            title="选择Excel文件",
            filetypes=[("Excel文件", "*.xlsx;*.xls")]
        )
        if file_path:
            self.excel_path.set(file_path)
            # 如果输出路径未设置，自动生成
            if not self.output_path.get():
                dir_path = os.path.dirname(file_path)
                file_name = os.path.basename(file_path)
                name_without_ext = os.path.splitext(file_name)[0]
                output_file = os.path.join(dir_path, f"{name_without_ext}_处理结果.docx")
                self.output_path.set(output_file)
            self._log(f"已选择Excel文件: {file_path}")
    
    def _browse_word(self):
        file_path = filedialog.askopenfilename(
            title="选择Word模板",
            filetypes=[("Word文件", "*.docx;*.doc")]
        )
        if file_path:
            self.word_path.set(file_path)
            # 如果输出路径未设置，自动生成
            if not self.output_path.get():
                dir_path = os.path.dirname(file_path)
                file_name = os.path.basename(file_path)
                name_without_ext = os.path.splitext(file_name)[0]
                output_file = os.path.join(dir_path, f"{name_without_ext}_处理结果.docx")
                self.output_path.set(output_file)
            self._log(f"已选择Word模板文件: {file_path}")
    
    def _browse_output(self):
        file_path = filedialog.asksaveasfilename(
            title="保存输出文件",
            defaultextension=".docx",
            filetypes=[("Word文件", "*.docx")]
        )
        if file_path:
            self.output_path.set(file_path)
            self._log(f"已设置输出文件路径: {file_path}")
    
    def _start_processing(self):
        # 检查参数
        if not self.excel_path.get():
            messagebox.showerror("错误", "请选择Excel文件")
            return
        
        if not self.word_path.get():
            messagebox.showerror("错误", "请选择Word模板")
            return
        
        if not self.output_path.get():
            messagebox.showerror("错误", "请指定输出文件路径")
            return
        
        try:
            copy_count = int(self.copy_count.get())
            if copy_count <= 0:
                raise ValueError
        except ValueError:
            messagebox.showerror("错误", "复制次数必须为正整数")
            return
        
        # 检查文件是否存在
        if not os.path.exists(self.excel_path.get()):
            messagebox.showerror("错误", f"Excel文件不存在: {self.excel_path.get()}")
            return
        
        if not os.path.exists(self.word_path.get()):
            messagebox.showerror("错误", f"Word模板不存在: {self.word_path.get()}")
            return
        
        # 检查输出文件是否被锁定
        output_dir = os.path.dirname(self.output_path.get())
        if not os.path.exists(output_dir):
            try:
                os.makedirs(output_dir)
            except Exception as e:
                messagebox.showerror("错误", f"无法创建输出目录: {str(e)}")
                return
        
        # 禁用按钮
        self.process_button.config(state=tk.DISABLED)
        self.cancel_button.config(state=tk.NORMAL)
        self.status_var.set("处理中...")
        self.is_processing = True
        
        # 清空日志
        self.log_text.config(state=tk.NORMAL)
        self.log_text.delete(1.0, tk.END)
        self.log_text.config(state=tk.DISABLED)
        
        # 重置进度条
        self.progress_var.set(0)
        self.progress_label.config(text="开始准备处理...")
        
        # 显示处理开始信息
        self._log("="*50)
        self._log("压实度文档处理开始")
        self._log(f"Excel文件: {self.excel_path.get()}")
        self._log(f"Word模板: {self.word_path.get()}")
        self._log(f"输出文件: {self.output_path.get()}")
        self._log(f"复制次数: {copy_count}")
        self._log("="*50)
        
        # 在新线程中处理，避免UI冻结
        self.process_thread = threading.Thread(
            target=self._process_files,
            args=(self.excel_path.get(),
                  self.word_path.get(),
                  copy_count,
                  self.output_path.get()),
            daemon=True
        )
        self.process_thread.start()
    
    def _cancel_processing(self):
        """取消处理过程"""
        if self.is_processing:
            self.is_processing = False
            self._log("用户取消了处理过程")
            self.root.after(100, lambda: self.progress_label.config(text="处理已取消"))
            
            # 显示取消消息
            messagebox.showinfo("已取消", "文档处理已被取消")
    
    def _process_files(self, excel_path, word_path, copy_count, output_path):
        try:
            # 确保使用绝对路径
            excel_path = os.path.abspath(excel_path)
            word_path = os.path.abspath(word_path)
            output_path = os.path.abspath(output_path)
            
            # 复制原文件的主要处理逻辑
            # 首先检查输出文件是否存在且被锁定
            if os.path.exists(output_path):
                try:
                    # 尝试以写模式打开文件，检查是否被锁定
                    with open(output_path, 'r+b'):
                        pass
                except PermissionError:
                    self._log(f"警告：文件 '{output_path}' 正在被另一个程序使用")
                    # 生成一个带时间戳的新文件名，避免文件锁定问题
                    timestamp = int(time.time())
                    output_dir = os.path.dirname(output_path)
                    output_name = os.path.basename(output_path)
                    name_without_ext = os.path.splitext(output_name)[0]
                    ext = os.path.splitext(output_name)[1]
                    output_path = os.path.join(output_dir, f"{name_without_ext}_{timestamp}{ext}")
                    self._log(f"将使用新的输出文件名: {output_path}")
                    # 更新界面上的输出路径
                    self.root.after(100, lambda: self.output_path.set(output_path))
            
            # 更新进度
            self.root.after(100, lambda: self.progress_var.set(10))
            self.root.after(100, lambda: self.progress_label.config(text="正在初始化处理环境..."))
            
            # 调用主处理函数
            self._log("开始执行Excel到Word的自动化处理...")
            self.root.after(100, lambda: self.progress_var.set(20))
            self.root.after(100, lambda: self.progress_label.config(text="正在处理Excel数据并生成表格..."))
            
            success = run_excel_to_word_automation(excel_path, word_path, copy_count, output_path, self._log)
            
            if success and self.is_processing:
                # 根据选项处理'g/cm3'单位
                if self.option_process_gcm3.get():
                    self.root.after(100, lambda: self.progress_var.set(60))
                    self.root.after(100, lambda: self.progress_label.config(text="正在处理g/cm³单位..."))
                    convert_g_cm3_to_superscript(output_path, self._log)
                
                # 根据选项统一处理所有附表标题的字体格式
                if self.option_unify_fonts.get():
                    self.root.after(100, lambda: self.progress_var.set(80))
                    self.root.after(100, lambda: self.progress_label.config(text="正在统一附表标题字体格式..."))
                    unify_all_schedule_headings_font(output_path, self._log)
                
                # 根据选项删除空数据行
                if self.option_delete_empty_rows.get():
                    self.root.after(100, lambda: self.progress_var.set(90))
                    self.root.after(100, lambda: self.progress_label.config(text="正在删除空数据行..."))
                    delete_rows_based_on_last_column(output_path, self._log)
                
                # 更新进度为完成
                self.root.after(100, lambda: self.progress_var.set(100))
                self.root.after(100, lambda: self.progress_label.config(text="处理完成"))
                
                self._log("="*50)
                self._log("脚本执行成功！")
                self._log(f"处理后的文件已保存至: {output_path}")
                
                # 完成后显示成功消息
                if self.is_processing:
                    self.root.after(100, lambda: messagebox.showinfo(
                        "成功", 
                        f"文档处理成功！\n\n输出文件:\n{output_path}\n\n处理选项:\n" +
                        f"- 处理g/cm³单位: {'是' if self.option_process_gcm3.get() else '否'}\n" +
                        f"- 统一附表标题字体: {'是' if self.option_unify_fonts.get() else '否'}\n" +
                        f"- 删除空数据行: {'是' if self.option_delete_empty_rows.get() else '否'}"
                    ))
            
        except Exception as e:
            self._log(f"脚本执行失败: {str(e)}")
            self._log(traceback.format_exc())
            if self.is_processing:
                self.root.after(100, lambda: messagebox.showerror("错误", f"处理失败: {str(e)}"))
        finally:
            # 恢复界面状态
            self.root.after(100, self._processing_complete)
    
    def _processing_complete(self):
        """处理完成后恢复界面状态"""
        self.process_button.config(state=tk.NORMAL)
        self.cancel_button.config(state=tk.DISABLED)
        self.status_var.set("就绪")
        self.is_processing = False
    
    def _log(self, message):
        """向日志区域添加消息，并在前面添加时间戳"""
        def append_log():
            timestamp = time.strftime("%H:%M:%S")
            self.log_text.config(state=tk.NORMAL)
            self.log_text.insert(tk.END, f"[{timestamp}] {message}\n")
            self.log_text.see(tk.END)  # 滚动到最后
            self.log_text.config(state=tk.DISABLED)
        
        # 在主线程中更新UI
        self.root.after(10, append_log)
    
    def get_app_dir(self):
        """获取应用程序所在目录，无论是直接运行还是打包为exe"""
        if getattr(sys, 'frozen', False):
            # 打包为exe的情况
            return os.path.dirname(os.path.abspath(sys.executable))
        else:
            # 直接运行Python脚本的情况
            return os.path.dirname(os.path.abspath(__file__))
    
    def set_working_directory(self):
        try:
            app_dir = self.get_app_dir()
            os.chdir(app_dir)
            self._log(f"已设置工作目录: {app_dir}")
        except:
            # 如果无法设置工作目录，就使用系统临时目录
            temp_dir = os.environ.get('TEMP', os.environ.get('TMP', '/tmp'))
            os.chdir(temp_dir)
            self._log(f"使用临时目录作为工作目录: {temp_dir}")

# 核心功能函数

def get_cell_display_value(cell):
    if cell.value is None:
        return ""  # 如果单元格为空，返回空字符串
    if isinstance(cell.value, (float, int)):
        # 检查小数位数
        if cell.number_format:
            if '.' in cell.number_format:
                decimal_places = cell.number_format.split('.')[1].count('0')
                return f"{cell.value:.{decimal_places}f}"  # 保持小数位数
        return str(cell.value)  # 直接返回数值
    if isinstance(cell.value, str):
        return cell.value  # 返回字符串
    # 处理日期类型
    if isinstance(cell.value, datetime):
        return cell.value.strftime("%Y-%m-%d")  # 格式化日期
    return str(cell.value)  # 对于其他类型，返回字符串形式


def run_excel_to_word_automation(excel_path, word_path, copy_count, new_word_path, log_status=None):
    """Excel到Word的自动化处理主函数"""
    if log_status is None:
        log_status = print
    
    try:
        # 确保使用绝对路径
        excel_path = os.path.abspath(excel_path)
        word_path = os.path.abspath(word_path)
        new_word_path = os.path.abspath(new_word_path)
        
        log_status("开始Excel到Word的自动化处理...")
        log_status(f"Excel文件路径: {excel_path}")
        log_status(f"Word模板路径: {word_path}")
        log_status(f"输出文件路径: {new_word_path}")
        log_status(f"复制次数: {copy_count}")
        
        # 打开Excel文件并读取数据
        wb = load_workbook(excel_path)
        ws = wb.active
        
        # 读取B列的数据（从第二行开始）
        data_list = []
        for row in ws.iter_rows(min_row=2, min_col=2, max_col=2, values_only=True):
            if row[0] is not None:
                data_list.append(row[0])
        
        log_status(f"从Excel中读取了 {len(data_list)} 条数据")
        
        # 打开Word文档
        doc = Document(word_path)
        log_status("已打开Word模板文档")
        
        # 查找"附表1"段落
        schedule1_para = None
        schedule1_table = None
        for i, para in enumerate(doc.paragraphs):
            if "附表1" in para.text and "压实度检测结果表" in para.text:
                schedule1_para = para
                log_status(f"找到附表1标题段落，索引: {i}, 内容: {para.text}")
                
                # 查找段落后的表格
                try:
                    para_elem = para._element
                    next_elem = para_elem.getnext()
                    while next_elem is not None:
                        if next_elem.tag.endswith('tbl'):
                            # 匹配对应的表格对象
                            for table in doc.tables:
                                if table._element is next_elem:
                                    schedule1_table = table
                                    log_status(f"找到附表1表格，包含 {len(schedule1_table.rows)} 行，{len(schedule1_table.columns)} 列")
                                    break
                            if schedule1_table:
                                break
                        next_elem = next_elem.getnext()
                except Exception as e:
                    log_status(f"通过XML路径查找表格出错: {e}")
                    
                break
        
        if not schedule1_para:
            log_status("错误：未找到附表1标题段落")
            return False
        
        if not schedule1_table:
            log_status("错误：未找到附表1表格")
            return False
        
        # 复制表格样式
        log_status("开始复制表格样式...")
        
        # 创建一个临时表格用于复制样式
        style_cells = []
        for row in schedule1_table.rows:
            row_styles = []
            for cell in row.cells:
                # 复制单元格样式
                cell_style = {
                    'font': cell.text and cell.paragraphs[0].runs[0].font if cell.text else None,
                    'alignment': cell.paragraphs[0].alignment if cell.paragraphs else None,
                    'vertical_alignment': cell.vertical_alignment if hasattr(cell, 'vertical_alignment') else None
                }
                row_styles.append(cell_style)
            style_cells.append(row_styles)
        
        log_status("表格样式复制完成")
        
        # 处理数据并创建新表格
        processed_data = []
        current_group = []
        max_rows_per_table = 25  # 每个表格最多25行
        
        for item in data_list:
            current_group.append(item)
            if len(current_group) >= max_rows_per_table:
                processed_data.append(current_group.copy())
                current_group.clear()
        
        if current_group:
            processed_data.append(current_group)
        
        log_status(f"数据分组完成，共分成 {len(processed_data)} 组")
        
        # 复制表格
        for i, group in enumerate(processed_data):
            if i >= copy_count:  # 限制复制次数
                break
            
            # 创建新的表格标题
            new_title = doc.add_paragraph()
            new_title.text = f"附表{i+1} 压实度检测结果表（{group[0]}）"
            
            # 复制原标题的格式
            if schedule1_para.runs:
                for run in schedule1_para.runs:
                    new_run = new_title.add_run()
                    new_run.font.name = run.font.name
                    new_run.font.size = run.font.size
                    new_run.font.bold = run.font.bold
                    new_run.font.italic = run.font.italic
            
            # 复制表格
            new_table = doc.add_table(rows=len(schedule1_table.rows), cols=len(schedule1_table.columns))
            
            # 设置表格列宽
            for col_idx in range(len(schedule1_table.columns)):
                new_table.columns[col_idx].width = schedule1_table.columns[col_idx].width
            
            # 填充数据和样式
            for row_idx in range(len(schedule1_table.rows)):
                for col_idx in range(len(schedule1_table.columns)):
                    # 复制文本
                    source_cell = schedule1_table.cell(row_idx, col_idx)
                    target_cell = new_table.cell(row_idx, col_idx)
                    
                    # 清空目标单元格
                    for para in target_cell.paragraphs:
                        para.clear()
                    
                    # 复制内容
                    for para in source_cell.paragraphs:
                        new_para = target_cell.add_paragraph()
                        for run in para.runs:
                            new_run = new_para.add_run(run.text)
                            new_run.font.name = run.font.name
                            new_run.font.size = run.font.size
                            new_run.font.bold = run.font.bold
                            new_run.font.italic = run.font.italic
            
            log_status(f"已创建并填充第 {i+1} 个表格")
        
        # 保存新文档
        try:
            # 先保存到临时文件
            temp_path = new_word_path + ".tmp"
            doc.save(temp_path)
            # 如果临时保存成功，再重命名为最终文件名
            if os.path.exists(new_word_path):
                os.remove(new_word_path)
            os.rename(temp_path, new_word_path)
            log_status(f"已保存处理后的文档: {new_word_path}")
        except Exception as e:
            log_status(f"保存文档时出错: {e}")
            # 尝试直接保存到最终文件
            try:
                doc.save(new_word_path)
                log_status(f"已通过备用方式保存处理后的文档: {new_word_path}")
            except Exception as e2:
                log_status(f"备用保存方式也失败: {e2}")
                return False
        
        # 验证生成的表格数量
        try:
            # 重新打开文档验证
            verify_doc = Document(new_word_path)
            table_count = 0
            for para in verify_doc.paragraphs:
                if "附表" in para.text and "压实度检测结果表" in para.text:
                    table_count += 1
            log_status(f"验证成功：生成了 {table_count} 个附表")
        except Exception as e:
            log_status(f"验证文档时出错: {e}")
        
        return True
    except Exception as e:
        log_status(f"Excel到Word自动化处理失败: {e}")
        log_status(traceback.format_exc())
        return False


def unify_all_schedule_headings_font(word_doc_path, log_status=None):
    """统一处理文档末尾"附表X 压实度检测结果表（YYYYY）"格式标题的字体：汉字设置宋体加粗小五，数字设置Times New Roman加粗小五"""
    if log_status is None:
        log_status = print
    
    try:
        import re
        from docx import Document
        from docx.shared import Pt
        
        log_status("开始统一处理文档末尾附表标题的字体格式...")
        
        # 打开Word文档
        doc = Document(word_doc_path)
        
        # 查找文档末尾的"附表X 压实度检测结果表（YYYYY）"格式标题
        # 先查找所有可能的段落，然后筛选出符合特定格式的
        schedule_paragraphs = []
        for para in doc.paragraphs:
            para_text = para.text.strip()
            # 使用正则表达式匹配"附表X 压实度检测结果表（YYYYY）"格式的标题
            if re.match(r'^附表\d+\s+压实度检测结果表\（.*\）$', para_text):
                schedule_paragraphs.append(para)
        
        log_status(f"找到 {len(schedule_paragraphs)} 个符合格式的附表标题段落")
        
        # 处理每个附表标题段落
        for para in schedule_paragraphs:
            try:
                # 清除段落中的所有内容
                original_text = para.text
                para.clear()
                
                # 逐字符处理，分开汉字和数字
                current_run = None
                current_is_chinese = None
                
                for char in original_text:
                    # 判断字符类型
                    if '\u4e00' <= char <= '\u9fff' or char in '（）':  # 汉字或括号
                        is_chinese = True
                    elif char.isdigit() or char == 'X':  # 数字或X
                        is_chinese = False
                    else:  # 其他字符（空格、标点等），跟随前一个字符的格式
                        is_chinese = current_is_chinese if current_is_chinese is not None else True
                    
                    # 如果字符类型改变，创建新的run
                    if current_run is None or is_chinese != current_is_chinese:
                        current_run = para.add_run(char)
                        current_is_chinese = is_chinese
                        
                        # 设置字体
                        if is_chinese:
                            current_run.font.name = "宋体"
                            # 确保中文字体正确应用
                            try:
                                current_run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
                            except:
                                pass
                        else:
                            current_run.font.name = "Times New Roman"
                        current_run.font.size = Pt(9)  # 小五
                        current_run.font.bold = True  # 加粗
                    else:
                        # 同一类型的字符，添加到当前run
                        current_run.text += char
                
            except Exception as e:
                log_status(f"处理标题段落时出错: {e}")
        
        # 保存修改后的文档
        doc.save(word_doc_path)
        log_status("所有附表标题字体格式已统一处理完成")
        
        return True
        
    except Exception as e:
        log_status(f"统一处理附表标题字体格式时出错: {e}")
        import traceback
        traceback.print_exc()
        return False


def convert_g_cm3_to_superscript(word_doc_path, log_status=None):
    """在文档中所有表格中查找'g/cm3'单位，并将其中的3改为上标"""
    if log_status is None:
        log_status = print
    
    try:
        from docx import Document
        
        log_status("开始处理文档末尾附表中的'g/cm3'单位...")
        
        # 打开Word文档
        doc = Document(word_doc_path)
        
        # 查找所有包含"附表"文本的段落及其对应的表格
        schedule_paragraphs = []
        for para_idx, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if "附表" in para_text:
                schedule_paragraphs.append((para_idx, para))
        
        if not schedule_paragraphs:
            log_status("未找到包含'附表'文本的段落")
            return True
        
        # 直接处理文档中的所有表格，不再限制只处理最后一个附表段落之后的表格
        # 这样可以确保所有包含'g/cm3'的表格都会被处理
        schedule_tables = doc.tables
        log_status(f"找到文档中的所有表格，共 {len(schedule_tables)} 个")
        
        # 如果没有表格，直接返回
        if not schedule_tables:
            log_status("文档中没有表格")
            return True
        
        # 处理每个附表表格中的'g/cm3'文本
        total_conversions = 0
        
        for table_idx, table in enumerate(schedule_tables):
            log_status(f"处理附表表格 {table_idx+1}/{len(schedule_tables)}")
            
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # 检查单元格中的所有段落
                    for para in cell.paragraphs:
                        # 检查段落中是否包含'g/cm3'
                        if 'g/cm3' in para.text:
                            log_status(f"在表格{table_idx+1}的单元格({row_idx+1},{cell_idx+1})中找到'g/cm3'")
                            
                            # 保存原始文本以进行日志记录
                            original_text = para.text
                            
                            # 清空段落
                            para.clear()
                            
                            # 分割文本，找到所有'g/cm3'出现的位置
                            parts = original_text.split('g/cm3')
                            
                            for i, part in enumerate(parts):
                                # 添加普通文本部分
                                if part:
                                    run = para.add_run(part)
                                    # 尝试保留原始格式
                                    try:
                                        # 这里简单设置为普通格式，可以根据需要调整
                                        run.font.name = "宋体" if any('\u4e00'-'\u9fff' in char for char in part) else "Times New Roman"
                                        run.font.size = None  # 使用默认大小
                                    except:
                                        pass
                                
                                # 如果不是最后一部分，添加'g/cm'加上上标的'3'
                                if i < len(parts) - 1:
                                    # 添加'g/cm'
                                    cm_run = para.add_run('g/cm')
                                    try:
                                        cm_run.font.name = "Times New Roman"
                                        cm_run.font.size = None
                                    except:
                                        pass
                                    
                                    # 添加上标的'3'
                                    superscript_run = para.add_run('3')
                                    try:
                                        superscript_run.font.name = "Times New Roman"
                                        superscript_run.font.size = None
                                        # 设置为上标
                                        superscript_run.font.superscript = True
                                    except:
                                        pass
                                
                                # 增加转换计数
                                if i < len(parts) - 1:
                                    total_conversions += 1
        
        log_status(f"已完成所有附表表格的处理，共转换 {total_conversions} 处'g/cm3'单位")
        
        # 保存修改后的文档
        doc.save(word_doc_path)
        log_status(f"文档已保存到: {word_doc_path}")
        
        return True
        
    except Exception as e:
        log_status(f"处理'g/cm3'单位时出错: {e}")
        import traceback
        traceback.print_exc()
        return False


def delete_rows_based_on_last_column(word_doc_path, log_status=None):
    """删除Excel表格中备注列（最后一列）为空的行"""
    if log_status is None:
        log_status = print
    
    try:
        from docx import Document
        import re
        
        log_status("开始删除备注列为空的行...")
        
        # 打开Word文档
        doc = Document(word_doc_path)
        
        # 查找所有表格
        tables = doc.tables
        log_status(f"找到 {len(tables)} 个表格")
        
        if not tables:
            log_status("文档中没有表格")
            return True
        
        total_rows_deleted = 0
        
        # 处理每个表格
        for table_idx, table in enumerate(tables):
            log_status(f"处理表格 {table_idx+1}/{len(tables)}")
            
            # 判断是否为数据表格（通常至少有5行）
            if len(table.rows) < 5:
                log_status(f"表格 {table_idx+1} 行数太少，跳过处理")
                continue
            
            # 从后往前遍历行，这样删除时不会影响前面的索引
            rows_to_delete = []
            
            # 从第三行开始（跳过表头和第一行数据）
            for i in range(len(table.rows) - 1, 2, -1):
                row = table.rows[i]
                
                # 检查最后一列单元格
                last_cell = row.cells[-1]
                cell_text = last_cell.text.strip()
                
                # 判断是否为空或0.0或#DIV/0!
                if not cell_text or cell_text == "0.0" or cell_text == "#DIV/0!" or re.match(r'^0(\.0+)?$', cell_text):
                    rows_to_delete.append(i)
            
            # 记录要删除的行数
            log_status(f"在表格 {table_idx+1} 中找到 {len(rows_to_delete)} 行需要删除")
            total_rows_deleted += len(rows_to_delete)
            
            # 从后往前删除行
            for i in reversed(rows_to_delete):
                if i < len(table.rows):
                    # 获取行元素
                    row_elem = table.rows[i]._element
                    # 获取父元素
                    tbl = row_elem.getparent()
                    # 删除行
                    tbl.remove(row_elem)
        
        log_status(f"总共删除了 {total_rows_deleted} 行数据")
        
        # 保存修改后的文档
        doc.save(word_doc_path)
        log_status(f"文档已保存到: {word_doc_path}")
        
        return True
        
    except Exception as e:
        log_status(f"删除空行时出错: {e}")
        import traceback
        traceback.print_exc()
        return False

# 导入datetime，因为在get_cell_display_value函数中使用
from datetime import datetime

# 主函数
if __name__ == "__main__":
    root = tk.Tk()
    app = ExcelToWordApp(root)
    root.mainloop()