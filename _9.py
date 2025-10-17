from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import qn
import sys
import traceback
import os
import re
from datetime import datetime

# 获取附表1的字体格式
def get_heading_format(doc, heading):
    for para in doc.paragraphs:
        if heading in para.text and para.runs:
            return para.runs[0].font.name, para.runs[0].font.size, para.runs[0].bold
    return None, None, None  # Return None if not found

def get_cell_display_value(cell):
    if cell.value is None:
        return ""  # 如果单元格为空，返回空字符串
    if isinstance(cell.value, (float, int)):
        # 检查小数位数
        if cell.number_format:
            if '.' in cell.number_format:
                decimal_places = cell.number_format.split('.')[1].count('0')
                return f"{cell.value:.{decimal_places}f}"  # 保持小数位数
        return str(cell.value)  # 直接返回数值
    if isinstance(cell.value, str):
        return cell.value  # 返回字符串
    # 处理日期类型
    if isinstance(cell.value, datetime):
        return cell.value.strftime("%Y-%m-%d")  # 格式化日期
    return str(cell.value)  # 对于其他类型，返回字符串形式

def copy_table_with_xml(source_table, target_table, log_status):
    """通过XML复制表格，确保格式完全一致"""
    try:
        log_status(f"开始XML复制表格，源表格行数: {len(source_table.rows)}, 目标表格行数: {len(target_table.rows)}")
        
        # 获取源表格的XML
        source_xml = source_table._element.xml
        
        # 解析源表格的XML结构
        from lxml import etree
        source_root = etree.fromstring(source_xml)
        
        # 获取目标表格的XML元素
        target_element = target_table._element
        
        # 复制整个表格的XML结构
        # 先清空目标表格
        for child in list(target_element):
            target_element.remove(child)
        
        # 复制源表格的所有子元素
        for child in source_root:
            target_element.append(child)
        
        # 复制源表格的属性
        for attr_name, attr_value in source_root.attrib.items():
            target_element.set(attr_name, attr_value)
        
        log_status("XML复制完成")
        
    except Exception as e:
        log_status(f"XML复制表格时出错: {e}")
        traceback.print_exc()

def copy_table_with_clone(source_table, target_table, log_status):
    """通过克隆复制表格"""
    try:
        log_status("开始克隆复制表格")
        
        # 获取源表格的XML
        source_xml = source_table._element.xml
        
        # 解析并克隆XML
        from lxml import etree
        source_root = etree.fromstring(source_xml)
        cloned_root = etree.fromstring(etree.tostring(source_root))
        
        # 替换目标表格的XML
        target_element = target_table._element
        parent = target_element.getparent()
        index = parent.index(target_element)
        
        # 删除原表格元素
        parent.remove(target_element)
        
        # 插入克隆的表格元素
        parent.insert(index, cloned_root)
        
        log_status("克隆复制完成")
        
    except Exception as e:
        log_status(f"克隆复制表格时出错: {e}")
        traceback.print_exc()

def copy_table_with_deep_copy(source_table, target_table, log_status):
    """深度复制表格"""
    try:
        log_status("开始深度复制表格")
        
        # 复制表格样式
        if source_table.style:
            target_table.style = source_table.style
        
        # 复制表格属性
        if hasattr(source_table, '_element') and hasattr(target_table, '_element'):
            # 复制整个表格属性
            if source_table._element.tblPr is not None:
                target_table._element.tblPr = source_table._element.tblPr
            
            # 复制表格网格
            source_tblGrid = source_table._element.xpath('w:tblGrid')
            if source_tblGrid:
                target_tblGrid = target_table._element.xpath('w:tblGrid')
                if not target_tblGrid:
                    target_table._element.append(source_tblGrid[0])
                else:
                    target_table._element.replace(target_tblGrid[0], source_tblGrid[0])
        
        # 复制行和单元格
        for row_idx in range(len(source_table.rows)):
            if row_idx < len(target_table.rows):
                source_row = source_table.rows[row_idx]
                target_row = target_table.rows[row_idx]
                
                # 复制行属性
                if hasattr(source_row, '_element') and hasattr(target_row, '_element'):
                    source_trPr = source_row._element.xpath('w:trPr')
                    if source_trPr:
                        target_trPr = target_row._element.xpath('w:trPr')
                        if not target_trPr:
                            target_row._element.insert(0, source_trPr[0])
                        else:
                            target_row._element.replace(target_trPr[0], source_trPr[0])
                
                # 复制单元格
                for col_idx in range(len(source_row.cells)):
                    if col_idx < len(target_row.cells):
                        source_cell = source_row.cells[col_idx]
                        target_cell = target_row.cells[col_idx]
                        
                        # 复制单元格属性
                        if hasattr(source_cell, '_element') and hasattr(target_cell, '_element'):
                            source_tcPr = source_cell._element.xpath('w:tcPr')
                            if source_tcPr:
                                target_tcPr = target_cell._element.xpath('w:tcPr')
                                if not target_tcPr:
                                    target_cell._element.insert(0, source_tcPr[0])
                                else:
                                    target_cell._element.replace(target_tcPr[0], source_tcPr[0])
                        
                        # 强制设置第三行的水平垂直居中对齐
                        if row_idx == 2:  # 第三行（索引为2）
                            if hasattr(target_cell, '_element'):
                                tcPr = target_cell._element.xpath('w:tcPr')
                                if tcPr:
                                    vAlign = parse_xml('<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>')
                                    tcPr[0].append(vAlign)
                                    textAlign = parse_xml('<w:textAlignment xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>')
                                    tcPr[0].append(textAlign)
                        
                        # 复制文本内容
                        target_cell.text = source_cell.text
        
        log_status("深度复制完成")
        
    except Exception as e:
        log_status(f"深度复制表格时出错: {e}")
        traceback.print_exc()

def should_refresh_via_excel(excel_path, table_ranges, col_range, all_data):
    # 判断除首列外是否全部为空
    for section in all_data:
        non_empty_beyond_first = any(any((str(val) if val is not None else "") != "" for val in row[1:]) for row in section)
        if non_empty_beyond_first:
            return False
    # 若所有分段除首列外都为空，再检查是否包含公式
    try:
        tmp_wb = load_workbook(excel_path, data_only=False)
        tmp_ws = tmp_wb.active
        for rng in table_ranges:
            for row in tmp_ws.iter_rows(min_row=rng[0], max_row=rng[1], min_col=col_range[0], max_col=col_range[1]):
                for cell in row:
                    val = cell.value
                    if isinstance(val, str) and val.startswith("="):
                        return True
        return False
    except Exception:
        return False

def refresh_excel_values_via_com(excel_file_path, log_status):
    try:
        import win32com.client  # 需要已安装 Excel 和 pywin32
    except Exception as e:
        log_status(f"提示：检测到可能是公式未计算导致的数据为空，但未安装 pywin32 或不可用，跳过自动刷新。错误: {e}")
        return False
    try:
        excel = win32com.client.Dispatch("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False
        wb = excel.Workbooks.Open(excel_file_path)
        try:
            wb.RefreshAll()
        except Exception:
            pass
        try:
            excel.CalculateFullRebuild()
        except Exception:
            try:
                excel.CalculateFull()
            except Exception:
                pass
        wb.Save()
        wb.Close(SaveChanges=False)
        excel.Quit()
        log_status("已通过 Excel 触发公式重算并保存。")
        return True
    except Exception as e:
        log_status(f"通过 Excel 触发重算失败: {e}")
        return False

def format_date(english_date_str):
    # 辅助函数：将如 "2025.1.4" 的格式转换为 "2025年1月4日"
    parts = english_date_str.split('.')
    if len(parts) == 3:
        year, month, day = parts
        return f"{year}年{month}月{day}日"
    return english_date_str  # 如果格式不匹配，返回原字符串

def process_remark_for_single_table(doc, ws, excel_path, target_table, original_section_index, actual_target_heading_text, log_status):
    # wb = load_workbook(excel_path, data_only=True) # 每次调用时重新加载
    # ws = wb.active
    
    row_base = 8  # 起始行号
    row_increment = 24  # 每隔 24 行递增
    s_row_base = 4  # S 列起始行号
    
    # 根据 original_section_index 计算对应的 Excel 行号
    current_row = row_base + original_section_index * row_increment  # 计算当前表格对应的行号
    s_cell_row = s_row_base + original_section_index * row_increment  # 计算 S 列的行号

    s_value_cell = f"S{s_cell_row}"  # 如 S4, S28 等
    c_value_cell = f"C{current_row}"  # 如 C8, C32 等
    k_value_cell = f"K{current_row}"  # 如 K8, K32 等
    c_value = get_cell_display_value(ws[c_value_cell])
    k_value = get_cell_display_value(ws[k_value_cell])
    
    # 查找第一列最后一行标题为"备注"的单元格
    remark_row = None
    for row in target_table.rows:
        if len(row.cells) > 0 and row.cells[0].text.strip() == "备注":
            remark_row = row
            break
    
    if not remark_row:
        log_status(f"错误：未找到 {actual_target_heading_text} 中的'备注'行")
        return
    
    if len(remark_row.cells) < 2:
        log_status(f"错误：{actual_target_heading_text} 的备注行没有第二列")
        return
    
    remark_text = remark_row.cells[1].text
    log_status(f"{actual_target_heading_text} 的原始备注文本: {remark_text}")
    
    new_remark_text = remark_text  # 初始化
    
    if "最大干密度：" in remark_text and "最佳含水率：" in remark_text:
        new_remark_text = remark_text.replace(
            "最大干密度：1.48g/cm3",
            f"最大干密度：{c_value}g/cm3"
        ).replace(
            "最佳含水率：14.4%",
            f"最佳含水率：{k_value}%"
        )
    
    # 执行日期替换逻辑
    s_value = get_cell_display_value(ws[s_value_cell])
    formatted_date = format_date(s_value)
    
    if "检测日期：2024年7月1日；检测方法：灌砂法" in remark_text:
        log_status(f"确认找到文本：{remark_text}")
        new_remark_text = new_remark_text.replace(
            "检测日期：2024年7月1日；检测方法：灌砂法",
            f"检测日期：{formatted_date}；检测方法：灌砂法"
        )
    else:
        log_status(f"警告：在 {actual_target_heading_text} 的备注文本中未找到指定的日期文本")
    
    remark_row.cells[1].text = new_remark_text
    log_status(f"{actual_target_heading_text} 的替换后备注文本: {new_remark_text}")

def group_sections_for_merging(all_sections_data, log_status):
    """根据比较值对相邻的表格部分进行分组以进行合并。"""
    grouped_sections = []
    if not all_sections_data: # 处理空数据情况
        return grouped_sections

    current_group = [all_sections_data[0]]

    for i in range(1, len(all_sections_data)):
        prev_section = all_sections_data[i-1]
        current_section = all_sections_data[i]

        # 比较所有指定的 Excel 单元格值
        should_merge = True
        for key in prev_section['comparison_values']:
            if prev_section['comparison_values'][key] != current_section['comparison_values'][key]:
                should_merge = False
                break
        
        if should_merge:
            log_status(f"发现相邻表格附表{prev_section['section_index']+1}和附表{current_section['section_index']+1}数据内容一致，进行合并。")
            current_group.append(current_section)
        else:
            grouped_sections.append(current_group)
            current_group = [current_section]
    
    grouped_sections.append(current_group) # 添加最后一组
    return grouped_sections

# 修改 run_excel_to_word_automation 函数以支持多个工作表
def run_excel_to_word_automation(excel_path, word_path, copy_count, new_word_path, status_callback=None):
    def log_status(message):
        if status_callback:
            status_callback(message)
        else:
            print(message)

    log_status("开始执行 Excel 到 Word 自动化，支持多个工作表。")
    log_status(f"Excel 文件路径: {excel_path}")
    log_status(f"Word 模板路径: {word_path}")
    log_status(f"复制次数: {copy_count}")
    log_status(f"输出 Word 文件路径: {new_word_path}")

    # 验证文件存在
    if not os.path.exists(excel_path):
        raise FileNotFoundError(f"错误：Excel文件不存在 - {excel_path}")
    if not os.path.exists(word_path):
        raise FileNotFoundError(f"错误：Word文件不存在 - {word_path}")

    try:
        wb = load_workbook(excel_path, data_only=True)  # 加载整个工作簿
        doc = Document(word_path)
        log_status(f"文档初始表格数量: {len(doc.tables)}")

        # 定义初始行和列范围
        start_row = 7
        end_row = 16
        col_range = (24, 29)  # X-AC列

        # 动态生成每个表格的数据范围和表头行数
        base_start_row = 7
        rows_per_section = 10 # 每组数据包含的行数 (例如 16 - 7 + 1 = 10)
        row_increment = 24 # 每组数据之间的行数增量

        # 新增：遍历所有工作表
        sheets = wb.sheetnames  # 获取所有工作表名称
        all_generated_tables_info = []  # 用于存储所有工作表的表格信息
        global_last_num = 1  # 全局附表编号，从1开始，确保连续

        # 定义表格数据和序号列的通用字体样式
        TABLE_DATA_FONT_NAME = "Times New Roman"
        TABLE_DATA_FONT_SIZE = Pt(10) # 5号字通常对应10磅
        TABLE_DATA_FONT_BOLD = False
  
        for sheet_name in sheets:
            ws = wb[sheet_name]  # 切换到当前工作表
            log_status(f"处理工作表: {sheet_name}")

            # 自动检测数据组数量，直到遇到停止条件或数据结束
            table_ranges = []
            header_rows_list = []
            i = 0
            while True:
                current_start = base_start_row + i * row_increment
                current_end = current_start + rows_per_section - 1
                
                if current_start > ws.max_row:  # 确保不超过当前工作表行数
                    break
                
                ac_cell = ws.cell(row=current_start, column=29)  # AC列是第29列
                ac_value = get_cell_display_value(ac_cell)
                
                if ac_value == "#DIV/0!":
                    log_status(f"检测到行{current_start} AC列为#DIV/0!，停止读取更多数据组")
                    break
                
                table_ranges.append((current_start, current_end))
                header_rows_list.append(1)  # 表头行数
                i += 1
                
                if i >= 100:
                    log_status("已达到最大数据组处理限制(100组)")
                    break

            # 读取B列和L列单元格内容并结合
            b_cell_values = []
            for i in range(len(table_ranges)):
                # 计算B列和L列的单元格引用
                b_cell_ref = f"B{5 + i*24}"  # B5, B29, B53, etc.
                l_cell_ref = f"L{3 + i*24}"  # L3, L27, L51, etc.
                
                # 读取两个单元格的值
                b_cell_value = get_cell_display_value(ws[b_cell_ref])
                l_cell_value = get_cell_display_value(ws[l_cell_ref])
                
                # 将B列和L列的值结合
                combined_value = b_cell_value + l_cell_value
                
                # 添加结合后的值到列表
                b_cell_values.append(combined_value)
                log_status(f"读取单元格 {b_cell_ref}={b_cell_value} 和 {l_cell_ref}={l_cell_value}，结合值={combined_value}")

            # 读取数据和合并所需的单元格内容
            all_sections_data = []
            actual_copy_count = len(table_ranges)  # 实际需要处理的数据组数
            
            for i in range(len(table_ranges)):
                current_start, current_end = table_ranges[i]
                log_status(f"正在读取附表{i+1}的Excel数据范围: 行{current_start}-{current_end}, 列{col_range[0]}-{col_range[1]}")
                
                # 检查当前组第一行AC列是否为#DIV/0!
                ac_cell = ws.cell(row=current_start, column=29)  # AC列是第29列
                ac_value = get_cell_display_value(ac_cell)
                
                if ac_value == "#DIV/0!":
                    log_status(f"检测到附表{i+1}第一行AC列为#DIV/0!，跳过该组及之后的数据组")
                    actual_copy_count = i  # 更新实际需要处理的数据组数
                    break  # 跳出循环，不再处理后续数据组
                
                section_data = []
                for row in ws.iter_rows(min_row=current_start, max_row=current_end, min_col=col_range[0], max_col=col_range[1]):
                    for cell in row:
                        log_status(f"单元格 {cell.coordinate} 的原始值: {cell.value}")  # 输出原始值
                    row_data = [get_cell_display_value(cell) for cell in row]  # 只获取数值
                    section_data.append(row_data)
                
                # 读取合并所需的额外单元格值
                comparison_values = {}
                row_offset = i * row_increment
                comparison_values['B3'] = get_cell_display_value(ws[f"B{3 + row_offset}"])
                comparison_values['B4'] = get_cell_display_value(ws[f"B{4 + row_offset}"])
                comparison_values['B5'] = get_cell_display_value(ws[f"B{5 + row_offset}"])
                comparison_values['L3'] = get_cell_display_value(ws[f"L{3 + row_offset}"])
                comparison_values['L5'] = get_cell_display_value(ws[f"L{5 + row_offset}"])
                comparison_values['P5'] = get_cell_display_value(ws[f"P{5 + row_offset}"])
                comparison_values['T5'] = get_cell_display_value(ws[f"T{5 + row_offset}"])
                comparison_values['S4'] = get_cell_display_value(ws[f"S{4 + row_offset}"])

                all_sections_data.append({
                    'section_index': i, # 原始的附表索引
                    'section_data': section_data,
                    'comparison_values': comparison_values,
                    'b_value': b_cell_values[i] # 原始的B列值，用于标题更新
                })
                log_status(f"附表{i+1}的数据: {section_data}")  # 打印读取的数据
            
            copy_count = len(table_ranges)  # 当前工作表的实际数据组数
            
            if copy_count == 0:
                continue  # 跳过空工作表

            # all_data 在此处被替换为 all_sections_data
            # 因此 should_refresh_via_excel 和 refresh_excel_values_via_com 需要调整
            # 为简化，暂时禁用应该重新加载excel的逻辑，直接使用all_sections_data
            # if should_refresh_via_excel(excel_path, table_ranges, col_range, all_sections_data):
            #     if refresh_excel_values_via_com(excel_path, log_status):
            #         wb = load_workbook(excel_path, data_only=True)
            #         ws = wb.active
            #         all_sections_data = [] # 需要重新读取所有数据
            #         for i in range(copy_count):
            #             # 重新读取 section_data 和 comparison_values
            #             pass # 此处省略重新读取逻辑，将在后续步骤中完善

            # 临时禁用重新加载Excel的逻辑，直接使用当前读取的all_sections_data

            # 根据合并规则对表格部分进行分组
            grouped_sections = group_sections_for_merging(all_sections_data, log_status)
            log_status(f"分组后的表格数量: {len(grouped_sections)}")

            # --- 步骤 1: 查找原始"附表1"段落和其下方的表格 (只执行一次) ---
            source_table = None
            first_heading_paragraph = None

            log_status("查找附表1段落...")
            for par_idx, para in enumerate(doc.paragraphs):
                if "附表1" in para.text and para.text.startswith("附表1"):
                    first_heading_paragraph = para
                    log_status(f"找到附表1段落，索引: {par_idx}, 内容: '{para.text}'")
                    
                    # 查找附表1下方的表格
                    log_status("查找附表1下方的表格...")
                    for table_idx, table in enumerate(doc.tables):
                        try:
                            # 检查表格元素是否紧跟在标题段落元素之后
                            if para._element.xpath('following-sibling::w:tbl[1]')[0] == table._element:
                                source_table = table
                                log_status(f"找到附表1下方的表格: 表格{table_idx+1}")
                                break
                        except Exception:
                            continue
                    
                    if source_table:
                        break
                    else:
                        log_status("未找到附表1下方的表格，尝试其他XML查找方法...")
                        try:
                            from lxml import etree
                            root = etree.fromstring(doc._element.xml)
                            para_elements = root.xpath(f"//w:p[contains(., '附表1')]")
                            if para_elements:
                                table_element = para_elements[0].xpath("following-sibling::w:tbl[1]")
                                if table_element:
                                    for table_idx, table in enumerate(doc.tables):
                                        if table._element == table_element[0]:
                                            source_table = table
                                            log_status(f"通过XML路径找到附表1下方的表格: 表格{table_idx+1}")
                                            break
                        except Exception as e:
                            log_status(f"XML查找方法出错: {e}")
                        
                        if source_table:
                            break
                        else:
                            log_status("错误：未找到附表1下方的表格")
                            raise Exception("未找到附表1下方的表格")

            if not first_heading_paragraph:
                log_status("错误：未找到附表1段落")
                raise Exception("未找到附表1段落")

            # 获取附表1的字体格式，用于后续新生成表格的标题格式
            first_heading_font_name, first_heading_font_size, first_heading_bold = get_heading_format(doc, "附表1")

            # 在循环之前，获取源表格的表头和备注行的索引
            source_header_row_idx = 0 # 假设表头是第一行
            source_remark_row_idx = -1
            for r_idx, row in enumerate(source_table.rows):
                if len(row.cells) > 0 and "备注" in row.cells[0].text:
                    source_remark_row_idx = r_idx
                    break
            
            # 获取源表格的列宽信息
            source_column_widths = []
            for col in source_table.columns:
                source_column_widths.append(col.width)
            log_status(f"源表格列宽: {source_column_widths}")

            # 提取源表格的表头、数据行和备注行XML元素作为模板
            header_row_xml_template = source_table.rows[source_header_row_idx]._element
            data_row_xml_template = source_table.rows[source_header_row_idx + 1]._element # 选择表头后的第一行作为数据行模板
            remark_row_xml_template = None
            if source_remark_row_idx != -1:
                remark_row_xml_template = source_table.rows[source_remark_row_idx]._element

            # 提取源表格的整体样式和网格信息
            source_tblPr_xml = source_table._element.tblPr
            source_tblGrid_xml = source_table._element.tblGrid

            data_rows_count = 12 # 每个表格的数据行数

            # 用于存储所有生成表格的列表，以便后续备注处理
            generated_tables_info = []

            # --- 步骤 2: 循环处理每个表格分组的数据填充和新表格生成 ---
            # for i in range(copy_count):
            for group_idx, current_group in enumerate(grouped_sections):
                # current_group 是一个列表，包含需要合并的 all_sections_data 字典
                # group_first_section 是当前组的第一个表格部分，用于获取标题、备注模板等
                group_first_section = current_group[0]
                
                log_status(f"\n=== 处理第{group_idx+1}个表格组 (包含 {len(current_group)} 个原始表格) ===")
                log_status(f"在 group_idx={group_idx} 循环开始时，文档表格数量: {len(doc.tables)}") # Debug: 打印循环开始时的表格数量
                
                current_target_table = None
                current_target_paragraph = None
                
                # 确定当前表格组的标题 (使用组内第一个表格的原始附表索引)
                # 附表编号将根据实际生成的表格顺序动态调整
                # new_num = group_first_section['section_index'] + 1 # 这个是原始附表编号，不能直接用作新附表编号
                # new_num 将在实际生成表格时计算

                # 如果是第一个表格组，直接填充原始的"附表1"表格
                if group_idx == 0 and sheet_name == sheets[0]: # 第一个工作表的第一个组
                    current_target_table = source_table
                    current_target_paragraph = first_heading_paragraph
                    target_heading_text = "附表1"
                else: # 新增附表标题和表格
                    # 获取最新附表编号
                    # last_num = 1 # 移除这行，因为它将被 global_last_num 替换
                    # 遍历 doc.paragraphs 来查找已有的附表编号，确保新编号是递增的
                    # for para in doc.paragraphs:
                    #     if para.text.startswith("附表"):
                    #         try:
                    #             num_part_candidate = ""
                    #             text_after_fubiao = para.text[2:].strip()
                    #             if text_after_fubiao and text_after_fubiao[0].isdigit():
                    #                 for char in text_after_fubiao:
                    #                     if char.isdigit():
                    #                         num_part_candidate += char
                    #                     else:
                    #                         break
                    #             
                    #             if num_part_candidate.isdigit():
                    #                 num = int(num_part_candidate)
                    #                 if num > last_num:
                    #                     last_num = num
                    #         except:
                    #             continue
                    
                    # new_num = last_num + 1 # 替换为使用 global_last_num
                    global_last_num += 1
                    new_num = global_last_num
                    target_heading_text = f"附表{new_num}"

                    # 添加新附表标题（复制附表1的格式）
                    doc.add_paragraph() # 先添加一个空行作为分隔
                    new_para = doc.add_paragraph()
                    new_run = new_para.add_run(f"附表{new_num} ") # 显式添加空格
                    # 先将new_para赋值给current_target_paragraph，然后再使用
                    current_target_paragraph = new_para
                    if new_num > 1: # 仅对附表2及之后的表格添加补充文本
                        # 保持完整的标题格式"压实度检测结果表（承台回填土）"
                        run_excel = current_target_paragraph.add_run("压实度检测结果表（承台回填土）")
                    # 使用从第一个附表获取的字体格式，如果获取失败则默认宋体9磅加粗
                    if first_heading_font_name and first_heading_font_size:
                        new_run.font.name = first_heading_font_name
                        new_run.font.size = first_heading_font_size
                        new_run.bold = first_heading_bold
                        if new_num > 1:
                            run_excel.font.name = first_heading_font_name
                            run_excel.font.size = first_heading_font_size
                            run_excel.bold = first_heading_bold
                    else:
                        new_run.font.name = "宋体"
                        new_run.font.size = Pt(9)
                        new_run.bold = True
                        if new_num > 1:
                            run_excel.font.name = "宋体"
                            run_excel.font.size = Pt(9)
                            run_excel.bold = True

                    new_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    current_target_paragraph = new_para

                    # 创建一个空的表格 (0行)
                    new_table = doc.add_table(rows=0, cols=len(source_table.columns))
                    tbl_element = new_table._element
                    from lxml import etree

                    # 复制表格整体属性
                    if source_tblPr_xml is not None:
                        tbl_element.insert(0, etree.fromstring(etree.tostring(source_tblPr_xml)))
                    if source_tblGrid_xml is not None:
                        tbl_element.append(etree.fromstring(etree.tostring(source_tblGrid_xml)))

                    # 设置新表格的列宽
                    if source_column_widths:
                        for col_idx, width in enumerate(source_column_widths):
                            if col_idx < len(new_table.columns):
                                new_table.columns[col_idx].width = width
                        log_status(f"新表格已设置列宽: {source_column_widths}")

                    # 插入克隆的表头行
                    tbl_element.append(etree.fromstring(etree.tostring(header_row_xml_template)))

                    # 插入数据行 (根据合并组的总数据行数)
                    total_data_rows_in_group = sum(len(section['section_data']) for section in current_group)
                    for _ in range(total_data_rows_in_group):
                        tbl_element.append(etree.fromstring(etree.tostring(data_row_xml_template)))

                    # 插入克隆的备注行 (如果存在)
                    if remark_row_xml_template is not None:
                        tbl_element.append(etree.fromstring(etree.tostring(remark_row_xml_template)))

                    current_target_table = new_table
                    log_status(f"在 group_idx={group_idx} 新增表格后，文档表格数量: {len(doc.tables)}") # Debug: 打印新增表格后的数量

                # 填充数据到当前目标表格
                # data_section = all_sections_data[i]['section_data'] # 替换为合并后的数据
                
                # 合并当前组的所有数据行
                merged_data_section = []
                for section in current_group:
                    merged_data_section.extend(section['section_data'])

                data_section_to_fill = merged_data_section
                header_rows = header_rows_list[0] # 表头行数统一使用第一个表格的
                
                if group_idx == 0: # 对于第一个表格，需要考虑其原始的行数，以及被删除的行
                    max_available_rows = len(current_target_table.rows) - header_rows - (1 if source_remark_row_idx != -1 else 0)
                    max_data_rows = min(len(data_section_to_fill), max_available_rows)
                else: # 对于后续生成的表格，行数已经精确控制
                    max_data_rows = len(data_section_to_fill) # 因为现在表格已经精确控制了行数，直接用实际数据行数
                
                log_status(f"正在填充 {target_heading_text}，数据行数: {len(data_section_to_fill)}，表格行数: {len(current_target_table.rows)}")
                for row_idx in range(max_data_rows):
                    row_data = data_section_to_fill[row_idx]

                    # 填充编号到第一列
                    num_cell = current_target_table.cell(row_idx + header_rows, 0)
                    # 清空现有段落并添加新运行以强制应用字体
                    for paragraph in list(num_cell.paragraphs):
                        num_cell._element.remove(paragraph._element)
                    p = num_cell.add_paragraph()
                    run = p.add_run(str(row_idx + 1))
                    # 应用表格数据和序号列的字体样式
                    run.font.name = TABLE_DATA_FONT_NAME
                    run.font.size = TABLE_DATA_FONT_SIZE
                    run.font.bold = TABLE_DATA_FONT_BOLD
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # 从Word表格的第二列开始粘贴，跳过第一列
                    for col_idx in range(min(len(row_data), len(current_target_table.columns) - 1)): # 减1因为跳过Word表格的第一列
                        cell = current_target_table.cell(row_idx + header_rows, col_idx + 1) # col_idx + 1 来从Word表格的第二列开始
                        cell_value = row_data[col_idx]
                        
                        # 清空现有段落并添加新运行以强制应用字体
                        for paragraph in list(cell.paragraphs):
                            cell._element.remove(paragraph._element)
                        p = cell.add_paragraph()
                        run = p.add_run(cell_value)
                        # 应用表格数据和序号列的字体样式
                        run.font.name = TABLE_DATA_FONT_NAME
                        run.font.size = TABLE_DATA_FONT_SIZE
                        run.font.bold = TABLE_DATA_FONT_BOLD
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # 确保垂直居中
                        from docx.oxml.shared import qn
                        tcPr = cell._element.xpath('w:tcPr')
                        if not tcPr:
                            tcPr = cell._element.makeelement(qn('w:tcPr'), nsmap=cell._element.nsmap)
                            cell._element.insert(0, tcPr)
                        else:
                            tcPr = tcPr[0]
                        
                        vAlign = tcPr.xpath('w:vAlign')
                        if not vAlign:
                            vAlign = tcPr.makeelement(qn('w:vAlign'), nsmap=tcPr.nsmap)
                            tcPr.append(vAlign)
                        vAlign = vAlign[0]  # Ensure vAlign is always set, even if it existed
                        vAlign.set(qn('w:val'), 'center')  # Explicitly set vertical centering

                # 在填充数据后，检查并删除行
                delete_rows_based_on_last_column(current_target_table, header_rows, log_status)

                # 移除二次强制设置所有数据单元格的垂直和水平居中的代码
                # for r_idx in range(header_rows, len(current_target_table.rows) - (1 if source_remark_row_idx != -1 else 0)):
                #     current_row = current_target_table.rows[r_idx]
                #     for cell in current_row.cells:
                #         # 确保水平居中
                #         for paragraph in cell.paragraphs:
                #             paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                #         
                #         # 确保垂直居中
                #         if hasattr(cell, '_element'):
                #             tcPr = cell._element.xpath('w:tcPr')
                #             if not tcPr:
                #                 tcPr = cell._element.makeelement(qn('w:tcPr'), nsmap=cell._element.nsmap)
                #                 cell._element.insert(0, tcPr)
                #             else:
                #                 tcPr = tcPr[0]
                #             
                #             vAlign = tcPr.xpath('w:vAlign')
                #             if not vAlign:
                #                 vAlign = tcPr.makeelement(qn('w:vAlign'), nsmap=tcPr.nsmap)
                #                 tcPr.append(vAlign)
                #             vAlign[0].set(qn('w:val'), 'top')  # 确保垂直居中改为靠上
                # log_status(f"已对 {target_heading_text} 的所有数据单元格强制应用垂直靠上水平居中。")

                # 更新标题段落中的Excel单元格文本 (如果适用)
                if current_target_paragraph:
                    try:
                        # 使用当前组的第一个表格部分的B列单元格值作为标题更新值
                        excel_value = group_first_section['b_value']
                        if excel_value and str(excel_value).strip():
                            # 只替换标题中的"承台回填土"部分，保留其他内容
                            if "压实度检测结果表（承台回填土）" in current_target_paragraph.text:
                                # 获取原始标题文本
                                original_text = current_target_paragraph.text
                                # 只替换"承台回填土"为Excel中的值
                                new_text = original_text.replace("承台回填土", str(excel_value))
                                # 清除现有内容
                                current_target_paragraph.clear()
                                # 添加更新后的文本
                                run_main = current_target_paragraph.add_run(new_text)
                            else:
                                # 如果标题格式不同，采用原有的清除重写方式
                                current_target_paragraph.clear()
                                run_main = current_target_paragraph.add_run(f"{target_heading_text} ")
                                run_excel = current_target_paragraph.add_run(str(excel_value))
                            # 应用从第一个附表标题获取的字体，如果获取失败则默认宋体9磅加粗
                            if first_heading_font_name and first_heading_font_size:
                                run_main.font.name = first_heading_font_name
                                run_main.font.size = first_heading_font_size
                                run_main.bold = first_heading_bold
                            else:
                                run_main.font.name = "宋体"
                                run_main.font.size = Pt(9)
                                run_main.bold = True
                            # 应用从第一个附表标题获取的字体，如果获取失败则默认宋体9磅加粗
                            if first_heading_font_name and first_heading_font_size:
                                run_excel.font.name = first_heading_font_name
                                run_excel.font.size = first_heading_font_size
                                run_excel.bold = first_heading_bold
                            else:
                                run_excel.font.name = "宋体"
                                run_excel.font.size = Pt(9)
                                run_excel.bold = True

                            current_target_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        log_status(f"更新 {target_heading_text} 标题时出错: {e}")
                else:
                    log_status(f"警告：未找到标题段落 '{target_heading_text}'，跳过更新其Excel单元格文本。")
                
                # 检查表格是否有数据行
                data_row_count = len(current_target_table.rows) - header_rows - (1 if source_remark_row_idx != -1 else 0)
                if data_row_count <= 0:
                    log_status(f"表格 {target_heading_text} 没有数据行，将删除整个表格")
                    # 删除表格
                    current_target_table._element.getparent().remove(current_target_table._element)
                    # 删除对应的标题段落
                    if current_target_paragraph:
                        current_target_paragraph._element.getparent().remove(current_target_paragraph._element)
                    continue
                
                # 记录当前生成的表格信息，用于后续的备注处理
                generated_tables_info.append({
                    'table_object': current_target_table,
                    'original_section_index': group_first_section['section_index'],
                    'target_heading_text': target_heading_text # 实际生成的附表标题
                })

            all_generated_tables_info.extend(generated_tables_info)  # 合并所有表格信息

        # 保存Word文件 (在所有表格生成和填充之后)
        # doc.save(new_word_path) # 移动到备注处理之后
        log_status(f"文档最终表格数量: {len(doc.tables)}") # Debug: 打印最终表格数量
        
        # 在脚本结束前，添加文档段落验证
        for p in doc.paragraphs:
            pass  # 如果需要，可以在这里添加其他处理逻辑

        # --- 步骤 3: 查找并替换"备注"行中的数值 ---
        # 原始 find_and_replace_remark 函数被替换为一个新的函数，以处理所有表格
        # process_all_remarks(doc, excel_path, copy_count, log_status)  # 替换为新的处理方式

        # 新的备注处理逻辑：遍历 generated_tables_info
        # 这里需要为 process_all_remarks 提供正确的 target_table 和 original_section_index
        # 注意：process_all_remarks 内部会重新加载 excel_path，所以这里只需要传递相关参数即可
        for table_info in all_generated_tables_info:
            target_table_obj = table_info['table_object']
            original_section_index = table_info['original_section_index']
            actual_target_heading_text = table_info['target_heading_text']
            
            # 调用 process_all_remarks，但是只处理单个表格的备注
            # process_all_remarks(doc, excel_path, 1, log_status) # copy_count 现在不适用
            
            # 需要一个修改后的 process_remark_for_single_table 函数
            process_remark_for_single_table(doc, ws, excel_path, target_table_obj, original_section_index, actual_target_heading_text, log_status)

        # 确保目录存在
        os.makedirs(os.path.dirname(new_word_path), exist_ok=True)
        
        # 保存文档
        temp_path = new_word_path + ".temp"
        doc.save(temp_path)
        log_status(f"文档已临时保存到: {temp_path}")
        
        # 验证文件是否成功生成
        if os.path.exists(temp_path):
            # 重命名为最终文件
            if os.path.exists(new_word_path):
                os.remove(new_word_path)
            os.rename(temp_path, new_word_path)
            log_status(f"文档已成功保存到: {new_word_path}")
            
            # 验证文件内容
            try:
                test_doc = Document(new_word_path)
                log_status(f"验证: 最终文档包含 {len(test_doc.tables)} 个表格")
            except Exception as e:
                log_status(f"验证文件时出错: {e}")
        else:
            log_status("错误: 文件保存失败")

        # 添加新逻辑：处理"表2 压实度检测结果评定表"表格的修改
        def modify_table2(doc, extracted_value, log_status, row_offset=0):
            try:
                # 1. 定位附表表格并计算压实度%平均值
                appendix_num = f"附表{row_offset + 1}"  # 动态生成附表编号
                appendix_table = None
                for para in doc.paragraphs:
                    if appendix_num in para.text:
                        # 查找段落后的第一个表格
                        for elem in para._element.xpath("following-sibling::*[position()<=3]"):
                            if elem.tag.endswith('tbl'):
                                for table in doc.tables:
                                    if table._element is elem:
                                        appendix_table = table
                                        break
                                if appendix_table:
                                    break
                        if appendix_table:
                            break
                
                if not appendix_table:
                    log_status(f"警告：未找到{appendix_num}表格")
                    return
                
                # 计算最后一列（压实度%）的平均值
                last_col_idx = len(appendix_table.columns) - 1
                values = []
                for row in appendix_table.rows[1:]:  # 跳过表头
                    if len(row.cells) > last_col_idx:
                        try:
                            val = float(row.cells[last_col_idx].text.strip('%'))
                            values.append(val)
                        except:
                            continue
                
                if not values:
                    log_status(f"警告：{appendix_num}最后一列没有有效数值")
                    return
                
                avg_value = round(sum(values) / len(values), 1)
                log_status(f"计算{appendix_num}压实度%平均值: {avg_value}%")
                
                # 2. 定位"表2 压实度检测结果评定表"后面的表格
                table2_found = False
                target_table = None
                
                # 更全面的表格标题匹配方式
                table_title_variants = [
                    "表2 压实度检测结果评定表",  # 原格式
                    "表2  压实度检测结果评定表",  # 两个空格
                    "表2压实度检测结果评定表",    # 无空格
                    "表2：压实度检测结果评定表",  # 冒号分隔
                    "表2.压实度检测结果评定表",   # 点号分隔
                    "表2-压实度检测结果评定表",   # 短横线分隔
                    "表2 压实度检测结果",         # 可能省略部分标题
                    "表2 压实度评定表"            # 更简化的标题
                ]
                
                # 遍历文档中的所有段落
                for para in doc.paragraphs:
                    # 检查所有可能的标题变体
                    para_text = para.text.strip()
                    for variant in table_title_variants:
                        if variant in para_text:
                            log_status(f"找到表格标题段落: '{para_text}'")
                            
                            # 方法1：使用XML路径查找紧邻的表格
                            try:
                                from lxml import etree
                                para_elem = para._element
                                # 查找段落后的第一个表格元素
                                next_elem = para_elem.getnext()
                                while next_elem is not None:
                                    if next_elem.tag.endswith('tbl'):
                                        # 匹配对应的表格对象
                                        for table in doc.tables:
                                            if table._element is next_elem:
                                                target_table = table
                                                table2_found = True
                                                log_status("通过XML路径找到表格")
                                                break
                                        if table2_found:
                                            break
                                    next_elem = next_elem.getnext()
                                if table2_found:
                                    break
                            except Exception as e:
                                log_status(f"XML路径查找出错: {e}")
                            
                            # 方法2：遍历所有表格，检查是否在段落附近
                            if not table2_found:
                                for table in doc.tables:
                                    table_elem = table._element
                                    # 检查表格是否在段落之后且距离不远
                                    if (para._element in table_elem.xpath("preceding-sibling::*") and 
                                        len(list(table_elem.xpath("preceding-sibling::*"))) - 
                                        len(list(para._element.xpath("following-sibling::*"))) <= 3):
                                        target_table = table
                                        table2_found = True
                                        log_status("通过相对位置找到表格")
                                        break
                            
                            if table2_found:
                                break
                    
                    if table2_found:
                        break
                
                if not table2_found:
                    log_status("警告：未找到 '表2 压实度检测结果评定表' 后面的表格")
                    return
                
                # 修改表格
                header_rows = 2  # 前两行是表头
                target_row_idx = 2 + row_offset  # 动态计算目标行索引
                
                # 检查是否需要添加新行
                while len(target_table.rows) <= target_row_idx:
                    # 添加新行（复制第三行的格式）
                    reference_row = target_table.rows[2]  # 第三行作为参考行
                    new_row = target_table.add_row()
                    
                    # 复制行高属性
                    if hasattr(reference_row, '_element'):
                        ref_trPr = reference_row._element.xpath('w:trPr')
                        if ref_trPr:
                            # 深拷贝trPr以确保所有属性（包括行高）被复制
                            cloned_trPr = etree.fromstring(etree.tostring(ref_trPr[0]))
                            new_trPr = new_row._element.xpath('w:trPr')
                            if not new_trPr:
                                new_row._element.insert(0, cloned_trPr)
                            else:
                                new_row._element.replace(new_trPr[0], cloned_trPr)
                    
                    # 复制单元格格式并移除垂直居中设置
                    for i in range(len(reference_row.cells)):
                        # 复制单元格属性
                        ref_cell = reference_row.cells[i]
                        new_cell = new_row.cells[i]
                        
                        # 复制单元格宽度和样式
                        ref_tcPr = ref_cell._element.xpath('w:tcPr')
                        if ref_tcPr:
                            new_tcPr = new_cell._element.xpath('w:tcPr')
                            if not new_tcPr:
                                new_cell._element.insert(0, etree.fromstring(etree.tostring(ref_tcPr[0]))) # 深拷贝 tcPr
                            else:
                                new_cell._element.replace(new_tcPr[0], etree.fromstring(etree.tostring(ref_tcPr[0])))
                        
                    log_status(f"已添加新行处理附表{row_offset + 1}的数据，并继承与第三行相同的行高和单元格样式")
                
                # 处理目标行
                row = target_table.rows[target_row_idx]
                
                # 第一列添加序号
                if len(row.cells) > 0:
                    num_cell = row.cells[0]
                    # 清空现有内容
                    for paragraph in list(num_cell.paragraphs):
                        num_cell._element.remove(paragraph._element)
                    # 添加新内容
                    p = num_cell.add_paragraph()
                    text_content = str(row_offset + 1)
                    # 逐字符设置字体
                    for char in text_content:
                        run = p.add_run(char)
                        if '\u4e00' <= char <= '\u9fff':  # 汉字
                            run.font.name = "宋体"
                        else:  # 其他字符
                            run.font.name = "Times New Roman"
                        run.font.size = Pt(10.5)  # 五号字
                        run.font.bold = False  # 常规
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 第二列放置提取的值
                    if len(row.cells) > 1:
                        value_cell = row.cells[1]
                        # 清空现有内容
                        for paragraph in list(value_cell.paragraphs):
                            value_cell._element.remove(paragraph._element)
                        # 添加新内容
                        p = value_cell.add_paragraph()
                        # 逐字符设置字体
                        for char in extracted_value:
                            run = p.add_run(char)
                            if '\u4e00' <= char <= '\u9fff':  # 汉字
                                run.font.name = "宋体"
                            else:  # 其他字符
                                run.font.name = "Times New Roman"
                            run.font.size = Pt(10.5)  # 五号字
                            run.font.bold = False  # 常规
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        log_status(f"已设置第二列格式，汉字使用宋体，非汉字使用Times New Roman")
                    
                    # 第三列放置设计值（≥94）
                    if len(row.cells) > 2:
                        design_cell = row.cells[2]
                        # 清空现有内容
                        for paragraph in list(design_cell.paragraphs):
                            design_cell._element.remove(paragraph._element)
                        # 添加新内容
                        p = design_cell.add_paragraph()
                        # 将≥符号和数字分开设置不同的字体
                        run_symbol = p.add_run("≥")  # ≥符号
                        run_symbol.font.name = "宋体"
                        run_symbol.font.size = Pt(10.5)  # 五号字
                        run_symbol.font.bold = False  # 常规
                        
                        run_number = p.add_run("94")  # 数字部分
                        run_number.font.name = "Times New Roman"
                        run_number.font.size = Pt(10.5)  # 五号字
                        run_number.font.bold = False  # 常规
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        log_status(f"已设置设计值格式: Times New Roman 10.5磅 常规")

                    # 第四列放置压实度%平均值
                    if len(row.cells) > 3:
                        avg_cell = row.cells[3]
                        # 清空现有内容
                        for paragraph in list(avg_cell.paragraphs):
                            avg_cell._element.remove(paragraph._element)
                        # 添加新内容
                        p = avg_cell.add_paragraph()
                        text_content = f"{avg_value}"
                        # 逐字符设置字体
                        for char in text_content:
                            run = p.add_run(char)
                            if '\u4e00' <= char <= '\u9fff':  # 汉字
                                run.font.name = "宋体"
                            else:  # 其他字符
                                run.font.name = "Times New Roman"
                            run.font.size = Pt(10.5)  # 五号字
                            run.font.bold = False  # 常规
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        log_status(f"已设置第四列格式，汉字使用宋体，非汉字使用Times New Roman")
                    
                    # 第五列和第六列填写数据行数
                    if len(row.cells) > 5:
                        # 计算附表表格的数据行数（除去首行和尾行）
                        data_row_count = max(0, len(appendix_table.rows) - 2)  # 减去首行和尾行
                        log_status(f"{appendix_num}表格数据行数: {data_row_count}")
                        
                        # 第五列（检测点数）
                        if len(row.cells) > 4:
                            count_cell = row.cells[4]
                            # 清空现有内容
                            for paragraph in list(count_cell.paragraphs):
                                count_cell._element.remove(paragraph._element)
                            # 添加新内容
                            p = count_cell.add_paragraph()
                            text_content = str(data_row_count)
                            # 逐字符设置字体
                            for char in text_content:
                                run = p.add_run(char)
                                if '\u4e00' <= char <= '\u9fff':  # 汉字
                                    run.font.name = "宋体"
                                else:  # 其他字符
                                    run.font.name = "Times New Roman"
                                run.font.size = Pt(10.5)  # 五号字
                                run.font.bold = False  # 常规
                            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            log_status(f"已设置检测点数: {data_row_count}")
                        
                        # 第六列（合格点数）
                        if len(row.cells) > 5:
                            qualified_cell = row.cells[5]
                            # 清空现有内容
                            for paragraph in list(qualified_cell.paragraphs):
                                qualified_cell._element.remove(paragraph._element)
                            # 添加新内容
                            p = qualified_cell.add_paragraph()
                            text_content = str(data_row_count)
                            # 逐字符设置字体
                            for char in text_content:
                                run = p.add_run(char)
                                if '\u4e00' <= char <= '\u9fff':  # 汉字
                                    run.font.name = "宋体"
                                else:  # 其他字符
                                    run.font.name = "Times New Roman"
                                run.font.size = Pt(10.5)  # 五号字
                                run.font.bold = False  # 常规
                            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            log_status(f"已设置合格点数: {data_row_count}")

                        # 第七列（合格率）
                        if len(row.cells) > 6:
                            rate_cell = row.cells[6]
                            # 清空现有内容
                            for paragraph in list(rate_cell.paragraphs):
                                rate_cell._element.remove(paragraph._element)
                            # 添加新内容
                            p = rate_cell.add_paragraph()
                            text_content = "100"
                            # 逐字符设置字体
                            for char in text_content:
                                run = p.add_run(char)
                                if '\u4e00' <= char <= '\u9fff':  # 汉字
                                    run.font.name = "宋体"
                                else:  # 其他字符
                                    run.font.name = "Times New Roman"
                                run.font.size = Pt(10.5)  # 五号字
                                run.font.bold = False  # 常规
                            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            log_status("已设置合格率: 100")

                        # 第八列（详细数据参见附表）
                        if len(row.cells) > 7:
                            ref_cell = row.cells[7]
                            # 清空现有内容
                            for paragraph in list(ref_cell.paragraphs):
                                ref_cell._element.remove(paragraph._element)
                            # 添加新内容
                            p = ref_cell.add_paragraph()
                            # 逐字符设置字体
                            full_text = f"附表{row_offset + 1}"
                            for char in full_text:
                                run = p.add_run(char)
                                if '\u4e00' <= char <= '\u9fff':  # 汉字
                                    run.font.name = "宋体"
                                else:  # 其他字符
                                    run.font.name = "Times New Roman"
                                run.font.size = Pt(10.5)  # 五号字
                                run.font.bold = False  # 常规
                            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            log_status(f"已设置详细数据参见附表: 附表{row_offset + 1}")
                
                log_status(f"成功将值'{extracted_value}'插入到压实度表格")

                # 处理目标行（第三行或根据 row_offset 计算的行）
                target_row_idx = 2 + row_offset  # 第三行索引
                if len(target_table.rows) > target_row_idx:
                    target_row = target_table.rows[target_row_idx]
                    
                    # 备用逻辑：如果上述设置无效，则复制原始表格的第三行格式
                    # 假设原始表格是 doc.tables[0]，您需要根据实际上下文调整
                    original_table = doc.tables[0]  # 假设这是原始的"表2"表格
                    if len(original_table.rows) > 2:
                        original_third_row = original_table.rows[2]
                        for cell_idx, cell in enumerate(target_row.cells):
                            if cell_idx < len(original_third_row.cells):
                                original_cell = original_third_row.cells[cell_idx]
                                # 复制单元格格式，包括字体、对齐等
                                # 注意：这里直接替换 tcPr 可能会导致其他格式丢失，需谨慎
                                # 更好的做法是逐个复制属性，或者在复制前保存并合并
                                if original_cell._element.find(qn('w:tcPr')) is not None:
                                    target_tcPr = cell._element.find(qn('w:tcPr'))
                                    if target_tcPr is not None:
                                        cell._element.replace(target_tcPr, original_cell._element.find(qn('w:tcPr')).copy())
                                    else:
                                        cell._element.insert(0, original_cell._element.find(qn('w:tcPr')).copy())
                                
                                # 清空并重新设置内容
                                for paragraph in list(cell.paragraphs): # 转换为列表以便修改时迭代
                                    cell._element.remove(paragraph._element)
                                p = cell.add_paragraph()
                                run = p.add_run(original_cell.text)
                                # 复制字体样式
                                if original_cell.paragraphs and original_cell.paragraphs[0].runs:
                                    original_run = original_cell.paragraphs[0].runs[0]
                                    run.font.name = original_run.font.name
                                    run.font.size = original_run.font.size
                                    run.font.bold = original_run.font.bold
                                p.paragraph_format.alignment = original_cell.paragraphs[0].paragraph_format.alignment
                    log_status("已应用备用逻辑：复制原始表格第三行格式")
                
            except Exception as e:
                log_status(f"修改表格时出错: {e}")
                traceback.print_exc()
        
        # 提取所有附表标题中的值
        extracted_values = []
        max_appendix_num = 0
        
        # 首先确定有多少个附表
        for para in doc.paragraphs:
            if "附表" in para.text:
                try:
                    # 提取附表编号
                    num_part = para.text.split("附表")[1].strip().split(" ")[0]
                    if num_part.isdigit():
                        appendix_num = int(num_part)
                        if appendix_num > max_appendix_num:
                            max_appendix_num = appendix_num
                except:
                    continue
        
        # 提取每个附表的值
        for i in range(1, max_appendix_num + 1):
            heading_text = f"附表{i}"
            value = extract_value_from_heading(doc, heading_text)
            if value:
                extracted_values.append((i-1, value))  # 存储(row_offset, value)
                log_status(f"从 '{heading_text}' 提取的值: {value}")
        
        if not extracted_values:
            log_status("警告：未从附表标题中提取到任何值")
            return
        
        # 保存修改前的表格内容用于验证
        pre_modify_tables = [(table, [cell.text for row in table.rows for cell in row.cells]) 
                           for table in doc.tables]
        
        # 打印修改前的表格内容
        log_status("修改前的表格内容:")
        for i, table in enumerate(doc.tables):
            log_status(f"表格{i+1}:")
            for row_idx, row in enumerate(table.rows):
                row_text = " | ".join(cell.text for cell in row.cells)
                log_status(f"行{row_idx}: {row_text}")
        
        # 在修改前保存文档状态
        temp_path_before = new_word_path + ".before"
        doc.save(temp_path_before)
        log_status(f"修改前文档已临时保存到: {temp_path_before}")
        
        # 执行表格修改，处理所有附表
        for row_offset, value in extracted_values:
            # if row_offset == 0: # 跳过附表1的数据处理到表2
            #     log_status(f"跳过附表{row_offset + 1}的数据处理到表2。")
            #     continue
            modify_table2(doc, value, log_status, row_offset=row_offset)
            
            # 立即保存修改后的文档
            temp_path_after = new_word_path + ".after"
            doc.save(temp_path_after)
            log_status(f"修改后文档已临时保存到: {temp_path_after}")
            
            # 验证修改是否成功
            target_table_modified = False
            log_status("修改后的表格内容:")
            for i, table in enumerate(doc.tables):
                log_status(f"表格{i+1}:")
                for row_idx, row in enumerate(table.rows):
                    row_text = " | ".join(cell.text for cell in row.cells)
                    log_status(f"行{row_idx}: {row_text}")
                    if row_idx >= 2 and len(row.cells) > 1:  # 检查所有数据行
                        cell_text = row.cells[1].text
                        for _, value in extracted_values:
                            if value and cell_text == value:
                                target_table_modified = True
                                log_status(f"验证成功: 表格{i+1}行{row_idx}第二列已更新为 '{cell_text}'")
                                break
                        else:
                            log_status(f"验证失败: 表格{i+1}行{row_idx}第二列值 '{cell_text}' 与预期值不匹配")
            
            if not target_table_modified and extracted_values:
                log_status("错误: 表格修改未生效，将尝试直接修改所有表格")
                # 尝试修改所有表格的对应行
                for i, table in enumerate(doc.tables):
                    for row_offset, value in extracted_values:
                        target_row = 2 + row_offset
                        if len(table.rows) > target_row and len(table.rows[target_row].cells) > 1:
                            cell = table.rows[target_row].cells[1]
                            # 清空现有内容
                            for paragraph in list(cell.paragraphs):
                                cell._element.remove(paragraph._element)
                            # 添加新内容
                            p = cell.add_paragraph()
                            run = p.add_run(value)
                            run.font.name = TABLE_DATA_FONT_NAME
                            run.font.size = TABLE_DATA_FONT_SIZE
                            run.font.bold = TABLE_DATA_FONT_BOLD
                            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            log_status(f"已强制修改表格{i+1}的行{target_row}第二列为 '{value}'")
            
            # 最终保存文档
            final_path = new_word_path
            if os.path.exists(temp_path_after):
                if os.path.exists(final_path):
                    os.remove(final_path)
                os.rename(temp_path_after, final_path)
                log_status(f"最终文档已保存到: {final_path}")
                
                # 在最终验证之前，修改表2对应的所有段落（从第3行开始，即段落50开始）
                log_status("开始修改所有表2对应段落...")
                try:
                    # 正确设置start_paragraph值为52，确保从段落52开始处理，也就是'本次对进行压实度检测...'这段话
                    modify_result = modify_all_paragraphs_from_table2_rows(final_path, start_paragraph=52, log_status=log_status)
                    if modify_result:
                        log_status("所有段落修改成功")
                    else:
                        log_status("所有段落修改失败")
                except Exception as e:
                    log_status(f"修改所有段落时出错: {e}")
                
                # 验证最终文件
                try:
                    test_doc = Document(final_path)
                    log_status(f"最终验证: 文档包含 {len(test_doc.tables)} 个表格")
                    for i, table in enumerate(test_doc.tables):
                        if len(table.rows) > 1 and len(table.rows[1].cells) > 1:
                            log_status(f"最终表格{i+1}第二行第二列值: '{table.rows[1].cells[1].text}'")
                    
                    # 验证所有表2对应的段落
                    log_status("验证所有表2对应段落...")
                    verified_count = 0
                    total_paragraphs = len(test_doc.paragraphs)
                    
                    # 检查从段落50开始的所有段落
                    for i in range(49, min(total_paragraphs, 60)):  # 最多检查到段落60
                        para_index = i
                        para_num = i + 1
                        para_text = test_doc.paragraphs[para_index].text
                        
                        # 检查是否包含压实度检测的内容
                        if "压实度检测" in para_text and "100%" in para_text:
                            log_status(f"段落{para_num}内容: '{para_text}'")
                            
                            # 检查100%是否加粗
                            bold_found = False
                            for run in test_doc.paragraphs[para_index].runs:
                                if "100%" in run.text and run.font.bold:
                                    bold_found = True
                                    log_status(f"段落{para_num}格式验证: 100%已设置为加粗")
                                    break
                            if not bold_found:
                                log_status(f"段落{para_num}格式警告: 100%可能未正确设置为加粗")
                            
                            verified_count += 1
                    
                    log_status(f"共验证了 {verified_count} 个表2对应段落")
                    
                except Exception as e:
                    log_status(f"最终验证出错: {e}")
            else:
                log_status("错误: 最终保存文件不存在")
        else:
            log_status("警告：未从 '附表1' 提取到值")

    except FileNotFoundError as e:
        log_status(str(e))
        raise
    except Exception as e:
        log_status(f"脚本执行失败: {e}")
        traceback.print_exc()
        raise

def delete_rows_based_on_last_column(table, header_rows, log_status):
    """根据最后一列的值删除表格行，保留备注行。"""
    try:
        log_status("开始检查并删除空/0.0的行...")

        remark_row_idx = -1
        for r_idx, row in enumerate(table.rows):
            if len(row.cells) > 0 and "备注" in row.cells[0].text:
                remark_row_idx = r_idx
                break

        if remark_row_idx == -1:
            log_status("警告：未找到备注行，无法确定删除范围。")
            return

        last_col_idx = len(table.columns) - 1
        rows_to_delete = []

        # 从备注行之前一行开始，从后往前遍历数据行
        # remark_row_idx 是备注行的索引，所以数据行范围是 header_rows 到 remark_row_idx - 1
        # 由于现在表格的行数已经控制，数据行范围就是从 header_rows 到 remark_row_idx - 1 (如果存在备注行) 或者 len(table.rows) - 1 (如果不存在备注行)
        
        # 确定数据行的上边界 (不包含备注行)
        upper_bound_data_rows = remark_row_idx if remark_row_idx != -1 else len(table.rows)

        for r_idx in range(upper_bound_data_rows - 1, header_rows - 1, -1):
            row = table.rows[r_idx]
            if last_col_idx < 0 or last_col_idx >= len(row.cells):
                log_status(f"警告：行 {r_idx} 的最后一列索引 {last_col_idx} 超出范围，跳过。")
                continue

            cell = row.cells[last_col_idx]
            cell_text = cell.text.strip()
            log_status(f"行 {r_idx} 最后一列原始文本: '{row.cells[last_col_idx].text}', strip后: '{cell_text}'") # 增加调试输出
            
            if cell_text == "" or cell_text == "0.0" or cell_text == "#DIV/0!":
                rows_to_delete.append(r_idx)
                log_status(f"标记删除行 {r_idx}，因为最后一列（'" + cell_text + "'）为空、0.0或#DIV/0!")

        # 实际删除行
        tbl = table._element
        for r_idx in sorted(rows_to_delete, reverse=True):
            tr = tbl.xpath(f'./w:tr[{r_idx + 1}]')[0] # XPath是1-based索引
            tbl.remove(tr)
            log_status(f"已删除行 {r_idx}")
            
        log_status("空/0.0行检查删除完成。")

    except Exception as e:
        log_status(f"删除空/0.0行时出错: {e}")
        traceback.print_exc()

# 新的 main 函数来兼容原始的直接运行方式，方便调试

def extract_value_from_heading(doc, heading_text):
    """从附表标题段落中提取值"""
    for para in doc.paragraphs:
        if heading_text in para.text:
            # 尝试提取标题中括号内的内容
            match = re.search(r'[（(](.*?)[）)]', para.text)
            if match:
                value = match.group(1).strip()
                if value:  # 如果提取到非空值
                    return value
            # 如果没有找到括号，回退到原来的逻辑
            parts = para.text.split(heading_text)
            if len(parts) > 1:
                value = parts[1].strip()
                if value:  # 如果提取到非空值
                    return value
    return None  # 如果没有找到值

def modify_paragraph_50_from_table2(word_doc_path, log_status=None):
    """从表2中提取数据并更新段落50的内容"""
    if log_status is None:
        log_status = print
    
    try:
        from docx import Document
        
        log_status("开始修改段落50，从表2提取数据...")
        
        # 打开Word文档
        doc = Document(word_doc_path)
        
        # 查找"表2"表格
        table2 = None
        table2_title_paragraph = None
        
        # 查找表2标题段落
        for para in doc.paragraphs:
            para_text = para.text.strip()
            # 尝试匹配多种可能的表2标题格式
            if any(variant in para_text for variant in [
                "表2 压实度检测结果评定表",
                "表2  压实度检测结果评定表",  # 两个空格
                "表2压实度检测结果评定表",    # 无空格
                "表2：压实度检测结果评定表",  # 冒号分隔
                "表2.压实度检测结果评定表",   # 点号分隔
                "表2-压实度检测结果评定表",   # 短横线分隔
                "表2 压实度检测结果",         # 可能省略部分标题
                "表2 压实度评定表"            # 更简化的标题
            ]):
                log_status(f"找到表2标题段落: '{para_text}'")
                table2_title_paragraph = para
                
                # 查找段落后的表格
                try:
                    from lxml import etree
                    para_elem = para._element
                    next_elem = para_elem.getnext()
                    while next_elem is not None:
                        if next_elem.tag.endswith('tbl'):
                            # 匹配对应的表格对象
                            for table in doc.tables:
                                if table._element is next_elem:
                                    table2 = table
                                    log_status(f"找到表2表格，包含 {len(table.rows)} 行，{len(table.columns)} 列")
                                    break
                            if table2:
                                break
                        next_elem = next_elem.getnext()
                    if table2:
                        break
                except Exception as e:
                    log_status(f"通过XML路径查找表2出错: {e}")
                    continue
        
        if not table2:
            log_status("错误：未找到'表2'表格")
            return False
        
        # 提取第三行（索引2）的数据
        if len(table2.rows) <= 2:
            log_status(f"错误：表2只有 {len(table2.rows)} 行，无法找到第三行")
            return False
        
        third_row = table2.rows[2]  # 第三行
        log_status(f"表2第三行内容: {[cell.text for cell in third_row.cells]}")
        
        # 提取第二列、第五列、第六列的数据
        if len(third_row.cells) < 6:
            log_status(f"错误：表2第三行只有 {len(third_row.cells)} 列，不足以提取需要的数据")
            return False
        
        second_col_value = third_row.cells[1].text.strip() if len(third_row.cells) > 1 else ""
        fifth_col_value = third_row.cells[4].text.strip() if len(third_row.cells) > 4 else ""
        sixth_col_value = third_row.cells[5].text.strip() if len(third_row.cells) > 5 else ""
        
        log_status(f"提取的数据 - 第二列: '{second_col_value}', 第五列: '{fifth_col_value}', 第六列: '{sixth_col_value}'")
        
        # 查找段落50
        if len(doc.paragraphs) < 50:
            log_status(f"错误：文档只有 {len(doc.paragraphs)} 个段落，无法找到段落50")
            return False
        
        paragraph_50 = doc.paragraphs[49]  # 段落50，索引为49
        original_text = paragraph_50.text
        log_status(f"段落50原始内容: '{original_text}'")
        
        # 检查段落50是否符合预期格式
        expected_text = "<w:rPr><w:b w:val=\"0\"/></w:rPr>（1）本次对进行压实度检测，检测点数为个，合格点数为个，合格率为100%。"
        if original_text != expected_text:
            log_status(f"警告：段落50内容与预期不符。实际: '{original_text}', 预期: '{expected_text}'")
        
        log_status(f"段落50原内容: '{original_text}'")
        
        # 更新段落50内容，分段设置格式（保持原格式但让100%加粗）
        paragraph_50.clear()
        
        # 创建多个run来实现不同的格式
        # 设置字体格式：宋体小四（12pt）用于中文，Times New Roman小四（12pt）用于数字
        chinese_font_name = "宋体"
        number_font_name = "Times New Roman"
        font_size = Pt(12)  # 小四号字体
        font_italic = False
        
        # 设置段落格式：首行缩进2字符，两端对齐
        paragraph_50.paragraph_format.first_line_indent = Pt(24)  # 2字符缩进（12pt * 2）
        paragraph_50.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐
        
        # 第一部分：本次对...进行压实度检测（中文部分）
        run1 = paragraph_50.add_run("本次对")
        run1.font.name = chinese_font_name
        run1.font.size = font_size
        run1.font.bold = False
        run1.font.italic = font_italic
        
        # 第二部分：数字部分（第二列值）
        run2 = paragraph_50.add_run(second_col_value)
        run2.font.name = number_font_name
        run2.font.size = font_size
        run2.font.bold = False
        run2.font.italic = font_italic
        
        # 第三部分：进行压实度检测，检测点数为（中文部分）
        run3 = paragraph_50.add_run("进行压实度检测，检测点数为")
        run3.font.name = chinese_font_name
        run3.font.size = font_size
        run3.font.bold = False
        run3.font.italic = font_italic
        
        # 第四部分：检测点数（第五列值）
        run4 = paragraph_50.add_run(fifth_col_value)
        run4.font.name = number_font_name
        run4.font.size = font_size
        run4.font.bold = False
        run4.font.italic = font_italic
        
        # 第五部分：个，合格点数为（中文部分）
        run5 = paragraph_50.add_run("个，合格点数为")
        run5.font.name = chinese_font_name
        run5.font.size = font_size
        run5.font.bold = False
        run5.font.italic = font_italic
        
        # 第六部分：合格点数（第六列值）
        run6 = paragraph_50.add_run(sixth_col_value)
        run6.font.name = number_font_name
        run6.font.size = font_size
        run6.font.bold = False
        run6.font.italic = font_italic
        
        # 第七部分：个，合格率为（中文部分）
        run7 = paragraph_50.add_run("个，合格率为")
        run7.font.name = chinese_font_name
        run7.font.size = font_size
        run7.font.bold = False
        run7.font.italic = font_italic
        
        # 第八部分：100%（加粗，使用Times New Roman字体）
        run8 = paragraph_50.add_run("100%")
        run8.font.name = number_font_name
        run8.font.size = font_size
        run8.font.bold = True   # 100%加粗
        run8.font.italic = font_italic
        
        # 第九部分：句号
        run9 = paragraph_50.add_run("。")
        run9.font.name = chinese_font_name
        run9.font.size = font_size
        run9.font.bold = False
        run9.font.italic = font_italic
        
        final_text = f"本次对{second_col_value}进行压实度检测，检测点数为{fifth_col_value}个，合格点数为{sixth_col_value}个，合格率为100%。"
        log_status(f"段落50新内容: '{final_text}' (其中100%为加粗)")
        
        # 保存文档
        doc.save(word_doc_path)
        log_status(f"段落50已更新并保存到: {word_doc_path}")
        
        return True
        
    except Exception as e:
        log_status(f"修改段落50时出错: {e}")
        import traceback
        traceback.print_exc()
        return False

def modify_paragraph_51_from_table2(word_doc_path, log_status=None):
    """从表2中提取第四行数据并更新段落51的内容"""
    if log_status is None:
        log_status = print
    
    try:
        from docx import Document
        
        log_status("开始修改段落51，从表2提取第四行数据...")
        
        # 打开Word文档
        doc = Document(word_doc_path)
        
        # 查找"表2"表格
        table2 = None
        table2_title_paragraph = None
        
        # 查找表2标题段落
        for para in doc.paragraphs:
            para_text = para.text.strip()
            # 尝试匹配多种可能的表2标题格式
            if any(variant in para_text for variant in [
                "表2 压实度检测结果评定表",
                "表2  压实度检测结果评定表",  # 两个空格
                "表2压实度检测结果评定表",    # 无空格
                "表2：压实度检测结果评定表",  # 冒号分隔
                "表2.压实度检测结果评定表",   # 点号分隔
                "表2-压实度检测结果评定表",   # 短横线分隔
                "表2 压实度检测结果",         # 可能省略部分标题
                "表2 压实度评定表"            # 更简化的标题
            ]):
                log_status(f"找到表2标题段落: '{para_text}'")
                table2_title_paragraph = para
                
                # 查找段落后的表格
                try:
                    from lxml import etree
                    para_elem = para._element
                    next_elem = para_elem.getnext()
                    while next_elem is not None:
                        if next_elem.tag.endswith('tbl'):
                            # 匹配对应的表格对象
                            for table in doc.tables:
                                if table._element is next_elem:
                                    table2 = table
                                    log_status(f"找到表2表格，包含 {len(table.rows)} 行，{len(table.columns)} 列")
                                    break
                            if table2:
                                break
                        next_elem = next_elem.getnext()
                    if table2:
                        break
                except Exception as e:
                    log_status(f"通过XML路径查找表2出错: {e}")
                    continue
        
        if not table2:
            log_status("错误：未找到'表2'表格")
            return False
        
        # 提取第四行（索引3）的数据
        if len(table2.rows) <= 3:
            log_status(f"错误：表2只有 {len(table2.rows)} 行，无法找到第四行")
            return False
        
        fourth_row = table2.rows[3]  # 第四行
        log_status(f"表2第四行内容: {[cell.text for cell in fourth_row.cells]}")
        
        # 提取第二列、第五列、第六列的数据
        if len(fourth_row.cells) < 6:
            log_status(f"错误：表2第四行只有 {len(fourth_row.cells)} 列，不足以提取需要的数据")
            return False
        
        second_col_value = fourth_row.cells[1].text.strip() if len(fourth_row.cells) > 1 else ""
        fifth_col_value = fourth_row.cells[4].text.strip() if len(fourth_row.cells) > 4 else ""
        sixth_col_value = fourth_row.cells[5].text.strip() if len(fourth_row.cells) > 5 else ""
        
        log_status(f"提取的数据 - 第二列: '{second_col_value}', 第五列: '{fifth_col_value}', 第六列: '{sixth_col_value}'")
        
        # 查找段落51
        if len(doc.paragraphs) < 51:
            log_status(f"错误：文档只有 {len(doc.paragraphs)} 个段落，无法找到段落51")
            return False
        
        paragraph_51 = doc.paragraphs[50]  # 段落51，索引为50
        original_text = paragraph_51.text
        log_status(f"段落51原始内容: '{original_text}'")
        
        # 检查段落51是否符合预期格式
        expected_text = "<w:rPr><w:b w:val=\"0\"/></w:rPr>（2）本次对进行压实度检测，检测点数为个，合格点数为个，合格率为100%。"
        if original_text != expected_text:
            log_status(f"警告：段落51内容与预期不符。实际: '{original_text}', 预期: '{expected_text}'")
        
        # 更新段落51内容，分段设置格式（保持原格式但让100%加粗）
        paragraph_51.clear()
        
        # 复制段落50的缩进与对齐
        try:
            template_paragraph = doc.paragraphs[49]
            paragraph_51.paragraph_format.left_indent = template_paragraph.paragraph_format.left_indent
            paragraph_51.paragraph_format.first_line_indent = template_paragraph.paragraph_format.first_line_indent
            paragraph_51.paragraph_format.alignment = template_paragraph.paragraph_format.alignment
        except Exception:
            pass
        
        # 创建多个run来实现不同的格式
        # 设置字体格式：宋体小四（12pt）用于中文，Times New Roman小四（12pt）用于数字
        chinese_font_name = "宋体"
        number_font_name = "Times New Roman"
        font_size = Pt(12)  # 小四号字体
        font_italic = False
        
        # 设置段落格式：首行缩进2字符，两端对齐
        paragraph_51.paragraph_format.first_line_indent = Pt(24)  # 2字符缩进（12pt * 2）
        paragraph_51.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐
        
        # 第一部分：本次对...进行压实度检测（中文部分）
        run1 = paragraph_51.add_run("本次对")
        run1.font.name = chinese_font_name
        run1.font.size = font_size
        run1.font.bold = False
        run1.font.italic = font_italic
        
        # 第二部分：数字部分（第二列值）
        run2 = paragraph_51.add_run(second_col_value)
        run2.font.name = number_font_name
        run2.font.size = font_size
        run2.font.bold = False
        run2.font.italic = font_italic
        
        # 第三部分：进行压实度检测，检测点数为（中文部分）
        run3 = paragraph_51.add_run("进行压实度检测，检测点数为")
        run3.font.name = chinese_font_name
        run3.font.size = font_size
        run3.font.bold = False
        run3.font.italic = font_italic
        
        # 第四部分：检测点数（第五列值）
        run4 = paragraph_51.add_run(fifth_col_value)
        run4.font.name = number_font_name
        run4.font.size = font_size
        run4.font.bold = False
        run4.font.italic = font_italic
        
        # 第五部分：个，合格点数为（中文部分）
        run5 = paragraph_51.add_run("个，合格点数为")
        run5.font.name = chinese_font_name
        run5.font.size = font_size
        run5.font.bold = False
        run5.font.italic = font_italic
        
        # 第六部分：合格点数（第六列值）
        run6 = paragraph_51.add_run(sixth_col_value)
        run6.font.name = number_font_name
        run6.font.size = font_size
        run6.font.bold = False
        run6.font.italic = font_italic
        
        # 第七部分：个，合格率为（中文部分）
        run7 = paragraph_51.add_run("个，合格率为")
        run7.font.name = chinese_font_name
        run7.font.size = font_size
        run7.font.bold = False
        run7.font.italic = font_italic
        
        # 第八部分：100%（加粗，使用Times New Roman字体）
        run8 = paragraph_51.add_run("100%")
        run8.font.name = number_font_name
        run8.font.size = font_size
        run8.font.bold = True   # 100%加粗
        run8.font.italic = font_italic
        
        # 第九部分：句号
        run9 = paragraph_51.add_run("。")
        run9.font.name = chinese_font_name
        run9.font.size = font_size
        run9.font.bold = False
        run9.font.italic = font_italic
        
        final_text = f"本次对{second_col_value}进行压实度检测，检测点数为{fifth_col_value}个，合格点数为{sixth_col_value}个，合格率为100%。"
        log_status(f"段落51新内容: '{final_text}' (其中100%为加粗)")
        
        # 保存文档
        doc.save(word_doc_path)
        log_status(f"段落51已更新并保存到: {word_doc_path}")
        
        return True
        
    except Exception as e:
        log_status(f"修改段落51时出错: {e}")
        import traceback
        traceback.print_exc()
        return False

# 查找独立附表标题整体部分
# 在执行段落50的复制新增操作期间，需要实现交叉运行机制
# 每当完成一个新增段落的复制插入后，立即根据该新增段落的行数，对"独立附表标题整体部分"执行向下移动操作
# 移动行数与新增段落的行数完全一致
# 移动方式通过在"独立附表标题（无数字）"的上一行连续按下两次Enter键实现换行

def modify_all_paragraphs_from_table2_rows(word_doc_path, start_paragraph=50, log_status=None):
    """处理表2所有后续行，从第4行开始，自动修改对应段落或复制段落50创建新段落"""
    if log_status is None:
        log_status = print
    
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Inches
        
        log_status(f"开始处理表2所有后续行，从第4行开始...")
        
        # 打开Word文档
        doc = Document(word_doc_path)
        
        # 查找独立附表标题（无数字标识）及其后续生成的所有附表内容
        # 这将作为需要下移的整体部分
        independent_schedule_title = None
        independent_schedule_content_start = None
        
        for para_idx, para in enumerate(doc.paragraphs):
            # 查找独立的"附表"标题（无数字标识）
            if "附表" in para.text and not any(char.isdigit() for char in para.text if char.strip()):
                log_status(f"找到独立附表标题，索引: {para_idx}, 内容: '{para.text}'")
                independent_schedule_title = para
                independent_schedule_content_start = para_idx
                break
        
        if not independent_schedule_title:
            log_status("警告：未找到独立附表标题，将跳过交叉运行机制")
            
        # 查找"表2"表格
        table2 = None
        log_status("查找表2表格...")
        
        # 查找表2标题段落
        for para in doc.paragraphs:
            para_text = para.text.strip()
            # 尝试匹配多种可能的表2标题格式
            if any(variant in para_text for variant in [
                "表2 压实度检测结果评定表",
                "表2  压实度检测结果评定表",  # 两个空格
                "表2压实度检测结果评定表",    # 无空格
                "表2：压实度检测结果评定表",  # 冒号分隔
                "表2.压实度检测结果评定表",   # 点号分隔
                "表2-压实度检测结果评定表",   # 短横线分隔
                "表2 压实度检测结果",         # 可能省略部分标题
                "表2 压实度评定表"            # 更简化的标题
            ]):
                log_status(f"找到表2标题段落: '{para_text}'")
                
                # 查找段落后的表格
                try:
                    from lxml import etree
                    para_elem = para._element
                    next_elem = para_elem.getnext()
                    while next_elem is not None:
                        if next_elem.tag.endswith('tbl'):
                            # 匹配对应的表格对象
                            for table in doc.tables:
                                if table._element is next_elem:
                                    table2 = table
                                    log_status(f"找到表2表格，包含 {len(table2.rows)} 行，{len(table2.columns)} 列")
                                    break
                            if table2:
                                break
                        next_elem = next_elem.getnext()
                    if table2:
                        break
                except Exception as e:
                    log_status(f"通过XML路径查找表2出错: {e}")
                    continue
        
        if not table2:
            log_status("错误：未找到'表2'表格")
            return False
        
        # 重写处理表2并生成新段落的逻辑
        def process_table2_and_generate_paragraphs(doc, table2, start_paragraph, independent_schedule_title=None, independent_schedule_content_start=None, log_status=None):
            """
            处理表2 压实度检测结果评定表，并根据表中数据生成新段落
            
            参数:
            doc: 文档对象
            table2: 表2对象
            start_paragraph: 开始处理的段落编号
            independent_schedule_title: 独立附表标题段落
            independent_schedule_content_start: 独立附表内容开始索引
            log_status: 日志记录函数
            
            返回:
            bool: 处理是否成功
            """
            if log_status is None:
                log_status = print
            
            if not table2:
                log_status("错误：未找到表2 压实度检测结果评定表")
                return False
            
            log_status(f"开始处理表2 压实度检测结果评定表，共{len(table2.rows)}行")
            
            # 从第3行开始处理（索引为2）
            start_row = 2
            
            # 检查是否有足够的段落作为模板
            if len(doc.paragraphs) < start_paragraph:
                log_status(f"错误：文档只有{len(doc.paragraphs)}个段落，无法找到段落{start_paragraph}作为模板")
                return False
            
            # 跟踪最后一个处理的段落索引
            last_paragraph_index = -1
            
            # 处理表2的数据行
            for row_index in range(start_row, len(table2.rows)):
                current_row = table2.rows[row_index]
                
                # 跳过列数不足的行
                if len(current_row.cells) < 6:
                    log_status(f"跳过第{row_index+1}行，列数不足")
                    continue
                
                # 检查第一列是否为数字（过滤非数据行）
                first_cell_text = current_row.cells[0].text.strip()
                if not first_cell_text or not first_cell_text.isdigit():
                    log_status(f"跳过第{row_index+1}行，第一列不是数字: '{first_cell_text}'")
                    continue
                
                # 提取需要的数据
                second_col_value = current_row.cells[1].text.strip() if len(current_row.cells) > 1 else ""
                fifth_col_value = current_row.cells[4].text.strip() if len(current_row.cells) > 4 else ""
                sixth_col_value = current_row.cells[5].text.strip() if len(current_row.cells) > 5 else ""
                
                log_status(f"\n处理表2第{row_index+1}行:")
                log_status(f"提取数据 - 第二列: '{second_col_value}', 第五列: '{fifth_col_value}', 第六列: '{sixth_col_value}'")
                
                # 计算目标段落索引（从start_paragraph开始）
                target_paragraph_index = start_paragraph - 1 + (row_index - start_row)
                target_paragraph_num = target_paragraph_index + 1
                
                # 获取或创建目标段落
                if len(doc.paragraphs) <= target_paragraph_index:
                    target_paragraph = doc.add_paragraph()
                    log_status(f"创建新段落{target_paragraph_num}")
                else:
                    target_paragraph = doc.paragraphs[target_paragraph_index]
                    # 清空现有内容
                    target_paragraph.clear()
                    log_status(f"准备修改段落{target_paragraph_num}")
                
                # 设置段落格式
                set_paragraph_format(target_paragraph, log_status)
                
                # 应用编号格式
                apply_numbering_format(target_paragraph, doc, start_paragraph, target_paragraph_num, log_status)
                
                # 构建并设置段落内容
                build_paragraph_content(target_paragraph, second_col_value, fifth_col_value, sixth_col_value, log_status)
                
                # 处理交叉运行机制（移动独立附表标题）
                if independent_schedule_title and target_paragraph_num > start_paragraph:
                    schedule_title_index = independent_schedule_content_start
                    # 估算新段落占用的行数
                    paragraph_text = f"本次对{second_col_value}进行压实度检测，检测点数为{fifth_col_value}个，合格点数为{sixth_col_value}个，合格率为100%。"
                    estimated_lines = max(1, int(len(paragraph_text) / 30) + 1)
                    
                    log_status(f"执行交叉运行机制：为段落{target_paragraph_num}在独立附表标题前添加{estimated_lines}个空行")
                    # 在独立附表标题前添加空段落以实现下移
                    for _ in range(estimated_lines):
                        empty_para = doc.add_paragraph()
                        empty_para.paragraph_format.line_spacing = 1.5
                        empty_para.paragraph_format.space_after = 0
                        # 将空段落移动到独立附表标题前
                        move_paragraph_before(empty_para, independent_schedule_title)
                    
                    # 更新独立附表标题引用
                    independent_schedule_title = doc.paragraphs[schedule_title_index + estimated_lines]
                    independent_schedule_content_start = schedule_title_index + estimated_lines
                
                # 更新最后处理的段落索引
                last_paragraph_index = target_paragraph_index
            
            # 在最后一个处理的段落之后添加日期段落
            if last_paragraph_index != -1:
                add_date_paragraph(doc, last_paragraph_index, log_status)
                # 关键：添加日期后终止换行移动操作
                independent_schedule_title = None
            
            # 检测并删除空白页
            remove_blank_pages(doc, log_status)
            
            # 将包含"附表"的标题移动到新页面顶部
            move_schedule_title_to_new_page(doc, log_status)
            
            # 保存文档
            doc.save(word_doc_path)
            log_status(f"所有段落已更新并保存到: {word_doc_path}")
            
            log_status("表2数据处理完成")
            return True
        
        def set_paragraph_format(paragraph, log_status):
            """设置段落的基础格式"""
            try:
                # 首行缩进2字符（小四字体约12pt，2字符即24pt）
                paragraph.paragraph_format.first_line_indent = Pt(24)
                # 两端对齐
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                log_status("已设置段落格式：首行缩进2字符，两端对齐")
            except Exception as e:
                log_status(f"设置段落格式时出错: {e}")
        
        def apply_numbering_format(paragraph, doc, start_paragraph, target_paragraph_num, log_status):
            """应用Word自动编号格式"""
            try:
                # 获取模板段落（通常是第start_paragraph个段落）
                template_para = doc.paragraphs[start_paragraph - 1]
                template_element = template_para._element
                
                # 查找编号属性
                template_num_pr = template_element.xpath('.//w:numPr')
                if not template_num_pr:
                    log_status(f"警告：段落{start_paragraph}没有找到编号格式模板")
                    return
                
                # 复制编号属性到目标段落
                target_element = paragraph._element
                
                # 确保目标段落有pPr元素
                target_p_pr = target_element.xpath('.//w:pPr')
                if not target_p_pr:
                    from docx.oxml import parse_xml
                    p_pr = parse_xml('<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                    target_element.insert(0, p_pr)
                    target_p_pr = target_element.xpath('.//w:pPr')
                
                if target_p_pr:
                    from lxml import etree
                    cloned_num_pr = etree.fromstring(etree.tostring(template_num_pr[0]))
                    target_p_pr[0].append(cloned_num_pr)
                    log_status(f"已为段落{target_paragraph_num}设置自动编号")
                    
                    # 确保编号字体为常规（不加粗）
                    ensure_normal_font_for_numbering(target_element, log_status)
                else:
                    log_status(f"警告：无法为段落{target_paragraph_num}设置编号格式")
            except Exception as e:
                log_status(f"设置段落{target_paragraph_num}编号格式时出错: {e}")
        
        def ensure_normal_font_for_numbering(element, log_status):
            """确保编号字体为常规样式（不加粗）"""
            try:
                # 移除编号相关的加粗设置
                num_font_elements = element.xpath('.//w:numPr//w:rPr//w:b')
                for font_elem in num_font_elements:
                    font_elem.getparent().remove(font_elem)
                
                # 为所有编号相关元素添加常规字体设置
                num_pr_elements = element.xpath('.//w:numPr')
                for num_pr in num_pr_elements:
                    r_pr = num_pr.xpath('.//w:rPr')
                    if not r_pr:
                        from docx.oxml import parse_xml
                        r_pr_elem = parse_xml('<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                        num_pr.append(r_pr_elem)
                        r_pr = [r_pr_elem]
                    
                    if r_pr:
                        # 移除现有加粗设置
                        bold_elems = r_pr[0].xpath('.//w:b')
                        for bold_elem in bold_elems:
                            r_pr[0].remove(bold_elem)
                        
                        # 添加常规字体设置
                        from docx.oxml import parse_xml
                        normal_font = parse_xml('<w:b w:val="false" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                        r_pr[0].append(normal_font)
                
                log_status("已确保编号字体为常规样式")
            except Exception as e:
                log_status(f"设置编号字体为常规时出错: {e}")
        
        def build_paragraph_content(paragraph, second_col_value, fifth_col_value, sixth_col_value, log_status):
            """构建段落内容，包括不同格式的文本部分"""
            # 字体设置
            chinese_font_name = "宋体"
            number_font_name = "Times New Roman"
            font_size = Pt(12)  # 小四号字体
            
            # 创建多个run来实现不同的格式
            # 第一部分：本次对...进行压实度检测（中文部分）
            run1 = paragraph.add_run("本次对")
            set_run_font(run1, chinese_font_name, font_size, bold=False)
            
            # 第二部分：数字部分（第二列值）
            run2 = paragraph.add_run(second_col_value)
            set_run_font(run2, number_font_name, font_size, bold=False)
            
            # 第三部分：进行压实度检测，检测点数为（中文部分）
            run3 = paragraph.add_run("进行压实度检测，检测点数为")
            set_run_font(run3, chinese_font_name, font_size, bold=False)
            
            # 第四部分：检测点数（第五列值）
            run4 = paragraph.add_run(fifth_col_value)
            set_run_font(run4, number_font_name, font_size, bold=False)
            
            # 第五部分：个，合格点数为（中文部分）
            run5 = paragraph.add_run("个，合格点数为")
            set_run_font(run5, chinese_font_name, font_size, bold=False)
            
            # 第六部分：合格点数（第六列值）
            run6 = paragraph.add_run(sixth_col_value)
            set_run_font(run6, number_font_name, font_size, bold=False)
            
            # 第七部分：个，合格率为（中文部分）
            run7 = paragraph.add_run("个，合格率为")
            set_run_font(run7, chinese_font_name, font_size, bold=False)
            
            # 第八部分：100%（加粗，使用Times New Roman字体）
            run8 = paragraph.add_run("100%")
            set_run_font(run8, number_font_name, font_size, bold=True)
            
            # 第九部分：句号
            run9 = paragraph.add_run("。")
            set_run_font(run9, chinese_font_name, font_size, bold=False)
            
            final_text = f"本次对{second_col_value}进行压实度检测，检测点数为{fifth_col_value}个，合格点数为{sixth_col_value}个，合格率为100%。"
            log_status(f"段落内容: '{final_text}' (其中100%为加粗)")
        
        def set_run_font(run, font_name, font_size, bold=False, italic=False):
            """设置文本运行的字体属性"""
            run.font.name = font_name
            run.font.size = font_size
            run.font.bold = bold
            run.font.italic = italic
            # 确保中文字体在所有语言设置中都正确应用
            if font_name in ["宋体", "黑体", "楷体"]:
                try:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
                except:
                    pass  # 忽略可能的异常
        
        def move_paragraph_before(source_para, target_para):
            """将源段落移动到目标段落之前"""
            try:
                source_elem = source_para._element
                target_elem = target_para._element
                
                # 从当前位置移除源段落
                source_parent = source_elem.getparent()
                if source_parent:
                    source_parent.remove(source_elem)
                    
                    # 在目标段落之前插入源段落
                    target_elem.addprevious(source_elem)
            except Exception as e:
                # 忽略可能的异常，保持程序运行
                pass
        
        def add_date_paragraph(doc, last_paragraph_index, log_status):
            """在最后一个处理的段落之后添加日期段落"""
            try:
                # 获取最后一个处理的段落
                last_paragraph = doc.paragraphs[last_paragraph_index]
                
                # 创建日期段落
                date_paragraph = doc.add_paragraph()
                
                # 设置段落格式：靠右对齐
                date_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # 添加日期文本
                date_text = "二○二五年六月十日"
                date_run = date_paragraph.add_run(date_text)
                
                # 设置字体格式：宋体小四加粗
                set_run_font(date_run, "宋体", Pt(12), bold=True)
                
                # 将日期段落移动到最后一个处理的段落之后
                move_paragraph_after(date_paragraph, last_paragraph)
                
                log_status(f"已添加日期段落: '{date_text}'")
            except Exception as e:
                log_status(f"添加日期段落时出错: {e}")
        
        def move_paragraph_after(source_para, target_para):
            """将源段落移动到目标段落之后"""
            try:
                source_elem = source_para._element
                target_elem = target_para._element
                
                # 从当前位置移除源段落
                source_parent = source_elem.getparent()
                if source_parent:
                    source_parent.remove(source_elem)
                    
                    # 在目标段落之后插入源段落
                    target_elem.addnext(source_elem)
            except Exception as e:
                # 忽略可能的异常，保持程序运行
                pass
        
        def remove_blank_pages(doc, log_status):
            """检测并删除文档中的空白页"""
            try:
                # 导入必要的模块
                import re
                from lxml import etree
                
                log_status("开始检测并删除文档中的空白页...")
                
                # 空白页检测和删除逻辑
                paragraphs_to_remove = []
                consecutive_empty_paragraphs = 0
                max_empty_paragraphs = 1  # 连续空白段落的阈值
                
                for i, para in enumerate(doc.paragraphs):
                    # 检查段落是否为空（没有实际内容）
                    is_empty_para = False
                    try:
                        # 检查段落中是否包含分页符
                        has_page_break = False
                        for run in para.runs:
                            run_element = run._element
                            br_elements = run_element.xpath('.//w:br[@w:type="page"]')
                            if br_elements:
                                has_page_break = True
                                break
                        
                        # 获取段落的实际文本内容（去除所有空白字符后的内容）
                        actual_text = re.sub(r'\s+', '', para.text)
                        
                        # 定义空白页段落的条件
                        if (has_page_break and actual_text == "") or actual_text == "":
                            is_empty_para = True
                    except Exception as ex:
                        log_status(f"检查段落{i}是否为空时出错: {ex}")
                        is_empty_para = False
                    
                    # 检测连续的空白段落
                    if is_empty_para:
                        consecutive_empty_paragraphs += 1
                        
                        # 如果发现连续的空白段落，且超过阈值，则标记为删除
                        if consecutive_empty_paragraphs > max_empty_paragraphs:
                            paragraphs_to_remove.append(i)
                            log_status(f"检测到空白页相关段落，标记段落{para.text!r} (索引{i}) 为删除")
                    else:
                        consecutive_empty_paragraphs = 0
                
                # 从后向前删除标记的段落，避免索引变化
                for i in reversed(paragraphs_to_remove):
                    if i < len(doc.paragraphs):
                        # 从文档中删除段落
                        p = doc.paragraphs[i]._element
                        p.getparent().remove(p)
                        p._p = p._element = None
                        log_status(f"已删除空白页相关段落，索引{i}")
                
                # 额外检查：如果文档末尾有空白页段落，删除它
                if doc.paragraphs:
                    # 检查最后几个段落是否都是空白的
                    end_empty_count = 0
                    for para in reversed(doc.paragraphs):
                        try:
                            actual_text = re.sub(r'\s+', '', para.text)
                            if actual_text == "":
                                end_empty_count += 1
                            else:
                                break
                        except:
                            break
                    
                    # 如果末尾有多个连续的空白段落，删除它们
                    if end_empty_count > 1:
                        for i in reversed(range(len(doc.paragraphs) - end_empty_count + 1, len(doc.paragraphs))):
                            if i < len(doc.paragraphs):
                                p = doc.paragraphs[i]._element
                                p.getparent().remove(p)
                                p._p = p._element = None
                                log_status(f"已删除文档末尾的空白页段落，索引{i}")
                
                log_status("空白页检测和删除完成")
                
            except Exception as e:
                log_status(f"检测和删除空白页时出错: {e}")
                import traceback
                traceback.print_exc()
        
        def move_schedule_title_to_new_page(doc, log_status):
            """查找包含'附表'文本但不包含数字的标题，并将其移动到新页面顶部"""
            try:
                # 导入必要的模块
                from docx.enum.text import WD_BREAK
                
                log_status("开始查找包含'附表'文本但不包含数字的标题，并将其移动到新页面顶部...")
                
                # 重新遍历文档段落，查找包含"附表"文本但不包含数字的标题
                target_para = None
                for para_idx, para in enumerate(doc.paragraphs):
                    para_text = para.text.strip()
                    # 检查是否包含"附表"文本且不包含任何数字
                    if "附表" in para_text and not any(char.isdigit() for char in para_text if char.strip()):
                        log_status(f"找到目标标题，索引: {para_idx}, 内容: '{para_text}'")
                        target_para = para
                        break
                
                if target_para:
                    # 在目标段落前插入分页符段落，确保标题移至新页面顶部
                    # 创建一个新的段落用于插入分页符
                    page_break_paragraph = doc.add_paragraph()
                    
                    # 在新段落中添加分页符
                    run = page_break_paragraph.add_run()
                    run.add_break(WD_BREAK.PAGE)
                    
                    # 将新段落移动到目标标题段落的前面
                    move_paragraph_before(page_break_paragraph, target_para)
                    
                    log_status(f"已在标题'附表'上方插入分页符段落，使其移动到新页面顶部")
                else:
                    log_status("未找到包含'附表'文本但不包含数字的标题")
                    
            except Exception as e:
                log_status(f"在移动标题到新页面顶部时出错: {e}")
                import traceback
                traceback.print_exc()
        
        # 执行重写后的处理函数
        try:
            return process_table2_and_generate_paragraphs(
                doc,
                table2,
                start_paragraph,
                independent_schedule_title,
                independent_schedule_content_start,
                log_status
            )
        except Exception as e:
            log_status(f"处理表2数据时出错: {e}")
            import traceback
            traceback.print_exc()
            return False

# 闭合 modify_all_paragraphs_from_table2_rows 函数
        return True
    except Exception as e:
        log_status(f"处理表2所有后续行时出错: {e}")
        import traceback
        traceback.print_exc()
        return False


def unify_all_schedule_headings_font(word_doc_path, log_status=None):
    """统一处理文档末尾"附表X 压实度检测结果表（YYYYY）"格式标题的字体：汉字设置宋体加粗小五，数字设置Times New Roman加粗小五"""
    if log_status is None:
        log_status = print
    
    try:
        import re
        from docx import Document
        from docx.shared import Pt
        
        log_status("开始统一处理文档末尾附表标题的字体格式...")
        
        # 打开Word文档
        doc = Document(word_doc_path)
        
        # 查找文档末尾的"附表X 压实度检测结果表（YYYYY）"格式标题
        # 先查找所有可能的段落，然后筛选出符合特定格式的
        schedule_paragraphs = []
        for para in doc.paragraphs:
            para_text = para.text.strip()
            # 使用正则表达式匹配"附表X 压实度检测结果表（YYYYY）"格式的标题
            if re.match(r'^附表\d+\s+压实度检测结果表\（.*\）$', para_text):
                schedule_paragraphs.append(para)
        
        log_status(f"找到 {len(schedule_paragraphs)} 个符合格式的附表标题段落")
        
        # 处理每个附表标题段落
        for para in schedule_paragraphs:
            try:
                # 清除段落中的所有内容
                original_text = para.text
                para.clear()
                
                # 逐字符处理，分开汉字和非汉字
                current_run = None
                current_is_chinese = None
                
                for char in original_text:
                    # 判断字符类型
                    if '\u4e00' <= char <= '\u9fff' or char in '（）':  # 汉字或括号
                        is_chinese = True
                    else:  # 所有其他字符（包括数字、字母、符号等）
                        is_chinese = False
                    
                    # 如果字符类型改变，创建新的run
                    if current_run is None or is_chinese != current_is_chinese:
                        current_run = para.add_run(char)
                        current_is_chinese = is_chinese
                        
                        # 设置字体
                        if is_chinese:
                            current_run.font.name = "宋体"
                        else:
                            current_run.font.name = "Times New Roman"
                        current_run.font.size = Pt(9)  # 小五
                        current_run.font.bold = True  # 加粗
                    else:
                        # 同一类型的字符，添加到当前run
                        current_run.text += char
                
            except Exception as e:
                log_status(f"处理标题段落时出错: {e}")
        
        # 保存修改后的文档
        doc.save(word_doc_path)
        log_status("所有附表标题字体格式已统一处理完成")
        
        return True
        
    except Exception as e:
        log_status(f"统一处理附表标题字体格式时出错: {e}")
        import traceback
        traceback.print_exc()
        return False

def convert_g_cm3_to_superscript(word_doc_path, log_status=None):
    """在文档中所有表格中查找'g/cm3'单位，并将其中的3改为上标"""
    if log_status is None:
        log_status = print
    
    try:
        from docx import Document
        from docx.enum.text import WD_BREAK
        
        log_status("开始处理文档末尾附表中的'g/cm3'单位...")
        
        # 打开Word文档
        doc = Document(word_doc_path)
        
        # 查找所有包含"附表"文本的段落及其对应的表格
        schedule_paragraphs = []
        for para_idx, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if "附表" in para_text:
                schedule_paragraphs.append((para_idx, para))
        
        if not schedule_paragraphs:
            log_status("未找到包含'附表'文本的段落")
            return True
        
        # 直接处理文档中的所有表格，不再限制只处理最后一个附表段落之后的表格
        # 这样可以确保所有包含'g/cm3'的表格都会被处理
        schedule_tables = doc.tables
        log_status(f"找到文档中的所有表格，共 {len(schedule_tables)} 个")
        
        # 如果没有表格，直接返回
        if not schedule_tables:
            log_status("文档中没有表格")
            return True
        
        # 处理每个附表表格中的'g/cm3'文本
        total_conversions = 0
        
        for table_idx, table in enumerate(schedule_tables):
            log_status(f"处理附表表格 {table_idx+1}/{len(schedule_tables)}")
            
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # 检查单元格中的所有段落
                    for para in cell.paragraphs:
                        # 检查段落中是否包含'g/cm3'
                        if 'g/cm3' in para.text:
                            log_status(f"在表格{table_idx+1}的单元格({row_idx+1},{cell_idx+1})中找到'g/cm3'")
                            
                            # 保存原始文本以进行日志记录
                            original_text = para.text
                            
                            # 清空段落
                            para.clear()
                            
                            # 分割文本，找到所有'g/cm3'出现的位置
                            parts = original_text.split('g/cm3')
                            
                            for i, part in enumerate(parts):
                                # 添加普通文本部分
                                if part:
                                    run = para.add_run(part)
                                    # 尝试保留原始格式
                                    try:
                                        # 这里简单设置为普通格式，可以根据需要调整
                                        run.font.name = "宋体" if any('\u4e00'-'\u9fff' in char for char in part) else "Times New Roman"
                                        run.font.size = None  # 使用默认大小
                                    except:
                                        pass
                                
                                # 如果不是最后一部分，添加'g/cm'加上上标的'3'
                                if i < len(parts) - 1:
                                    # 添加'g/cm'
                                    cm_run = para.add_run('g/cm')
                                    try:
                                        cm_run.font.name = "Times New Roman"
                                        cm_run.font.size = None
                                    except:
                                        pass
                                    
                                    # 添加上标的'3'
                                    superscript_run = para.add_run('3')
                                    try:
                                        superscript_run.font.name = "Times New Roman"
                                        superscript_run.font.size = None
                                        # 设置为上标
                                        superscript_run.font.superscript = True
                                    except:
                                        pass
                                
                                # 增加转换计数
                                if i < len(parts) - 1:
                                    total_conversions += 1
        
        log_status(f"已完成所有附表表格的处理，共转换 {total_conversions} 处'g/cm3'单位")
        
        # 保存修改后的文档
        doc.save(word_doc_path)
        log_status(f"文档已保存到: {word_doc_path}")
        
        return True
        
    except Exception as e:
        log_status(f"处理'g/cm3'单位时出错: {e}")
        import traceback
        traceback.print_exc()
        return False

# 注意：压实度标准值的修改已移至modify_table2函数中的相关代码，
# 在函数中分别设置了≥符号（宋体）和数字（Times New Roman）的格式

import time
import os

if __name__ == "__main__":
    # 源文件路径和目标文件路径
    excel_path = r"C:\Users\xc\Desktop\模版\路面路基模板\3.xlsx"
    word_path = r"C:\Users\xc\Desktop\模版\路面路基模板\4.docx"
    new_word_path = r"C:\Users\xc\Desktop\模版\路面路基模板\9_new.docx" # 新文件名
    copy_count =50 # 你可以根据需要调整复制次数
    
    # 检查目标文件是否存在且被锁定
    if os.path.exists(new_word_path):
        try:
            # 尝试以写模式打开文件，检查是否被锁定
            with open(new_word_path, 'r+b'):
                pass
        except PermissionError:
            print(f"警告：文件 '{new_word_path}' 正在被另一个程序使用，请关闭该文件后再运行脚本。")
            # 生成一个带时间戳的新文件名，避免文件锁定问题
            timestamp = int(time.time())
            new_word_path = fr"C:\Users\xc\Desktop\模版\路面路基模板\9_new_{timestamp}.docx"
            print(f"将使用新的输出文件名: {new_word_path}")
    
    try:
        run_excel_to_word_automation(excel_path, word_path, copy_count, new_word_path)
        # 在保存文档后调用处理'g/cm3'单位的函数
        convert_g_cm3_to_superscript(new_word_path)
        # 统一处理所有附表标题的字体格式
        unify_all_schedule_headings_font(new_word_path)
        print("脚本执行成功！")
        print(f"处理后的文件已保存至: {new_word_path}")
    except Exception as e:
        print(f"脚本执行失败: {e}")
        traceback.print_exc()
